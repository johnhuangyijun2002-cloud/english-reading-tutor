import asyncio
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Optional
from urllib.parse import urlencode

import asyncpg
import feedparser
from cryptography.fernet import Fernet, InvalidToken
import httpx
import jieba
import jieba.posseg as pseg
from janome.tokenizer import Tokenizer as JanomeTokenizer
import pdfplumber
import trafilatura
from docx import Document as DocxDocument
from dotenv import load_dotenv
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent.parent
# 本地开发默认用项目里的 data/ 文件夹；部署到 Railway 之类的平台时，把 DATA_DIR
# 这个环境变量指到挂载的持久化磁盘（比如 /data）。这些 JSON 文件现在只在启动时
# 用来做"迁移到 Postgres"的一次性数据源，迁移完之后正常读写都走数据库了。
DATA_DIR = Path(os.environ.get("DATA_DIR") or (BASE_DIR / "data"))
DOCUMENTS_FILE = DATA_DIR / "documents.json"
VOCAB_FILE = DATA_DIR / "vocab.json"
SENTENCE_NOTES_FILE = DATA_DIR / "sentence_notes.json"
USAGE_FILE = DATA_DIR / "api_usage.json"
USERS_FILE = DATA_DIR / "users.json"
SESSIONS_FILE = DATA_DIR / "sessions.json"
FRONTEND_DIR = BASE_DIR / "frontend"

DATA_DIR.mkdir(parents=True, exist_ok=True)

SHEETS_WEBHOOK_URL = os.environ.get("SHEETS_WEBHOOK_URL", "")
SHEETS_TOKEN = os.environ.get("SHEETS_TOKEN", "")

GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")

TURNSTILE_SITE_KEY = os.environ.get("TURNSTILE_SITE_KEY", "")
TURNSTILE_SECRET_KEY = os.environ.get("TURNSTILE_SECRET_KEY", "")

# 公共体验额度：没填自己 AI Key 的账户，可以先用站长自己出钱的 key 免费调用几次，
# 用完了再提示自己填 key。HOUSE_AI_API_KEY 没配置的话这个功能就是关闭状态，
# 完全不影响"必须自己填 key"这个原有行为。
HOUSE_AI_PROVIDER = os.environ.get("HOUSE_AI_PROVIDER", "deepseek")
HOUSE_AI_API_KEY = os.environ.get("HOUSE_AI_API_KEY", "")
HOUSE_FREE_CALLS_PER_USER = 10
HOUSE_MONTHLY_BUDGET_USD = float(os.environ.get("HOUSE_MONTHLY_BUDGET_USD", "5"))

DATABASE_URL = os.environ.get("DATABASE_URL", "")

ENCRYPTION_KEY = os.environ.get("ENCRYPTION_KEY", "")
if not ENCRYPTION_KEY:
    raise RuntimeError(
        "没有配置 ENCRYPTION_KEY，没法加密存储用户的 AI API Key。"
        "用 `python -c \"from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())\"` 生成一个，"
        "配到环境变量里(本地 .env / Railway variables)。"
    )
_fernet = Fernet(ENCRYPTION_KEY.encode())


def _encrypt_key(value: str) -> str:
    if not value:
        return value
    return _fernet.encrypt(value.encode()).decode()


def _decrypt_key(value: str) -> str:
    if not value:
        return ""
    try:
        return _fernet.decrypt(value.encode()).decode()
    except InvalidToken:
        return ""


def _is_encrypted(value: str) -> bool:
    if not value:
        return True
    try:
        _fernet.decrypt(value.encode())
        return True
    except InvalidToken:
        return False

app = FastAPI(title="English Reader")

# 限流：按客户端 IP 算(uvicorn 已经配了 --proxy-headers，Railway 边缘代理转发的真实 IP
# 能正确识别到)。主要防的是注册/登录被脚本刷，以及 AI 解析接口被刷导致 API 账单暴涨。
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def no_store_api_responses(request, call_next):
    # /api/* 响应带的是登录凭证鉴权后的个人数据，不能被浏览器/中间代理按 URL 缓存，
    # 不然换账号登录或者退出登录之后，可能还读到上一个人登录时缓存下来的响应。
    response = await call_next(request)
    if request.url.path.startswith("/api/"):
        response.headers["Cache-Control"] = "no-store"
    elif request.url.path in ("/", "/style.css", "/app.js") or request.url.path.startswith("/i18n/"):
        # 前端是纯静态文件、没走构建工具，index.html/style.css/app.js 之前完全没设缓存策略，
        # 浏览器会按启发式规则长期缓存——出现过好几次"部署了新版本，但用户浏览器还在跑
        # 旧的 app.js/style.css，导致新功能点了没反应、页面样式没更新"的问题。
        # no-cache 不是不缓存，是每次都问一下服务器"内容变了吗"(带 ETag 的条件请求，
        # 没变就是很快的 304)，保证只要部署了新版本，下一次打开就一定是新版本。
        response.headers["Cache-Control"] = "no-cache"
    return response


def _read_json(path: Path, default):
    if not path.exists():
        return default
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def _write_json(path: Path, data):
    tmp_path = path.with_name(path.name + f".tmp{os.getpid()}")
    with tmp_path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    tmp_path.replace(path)


# ---------- 数据库(Postgres/Supabase) ----------
# 账号密码、文档、生词、句子笔记、AI 用量都存在 Postgres 里，不再落地成 JSON 文件。
# 第一次启动、且数据库还是空的时候，会自动把 data/*.json 里的旧数据搬进去(见 init_db)。

_pool: Optional[asyncpg.Pool] = None

SCHEMA_SQL = """
CREATE TABLE IF NOT EXISTS users (
    id TEXT PRIMARY KEY,
    name TEXT NOT NULL DEFAULT '',
    username TEXT UNIQUE,
    password_salt TEXT,
    password_hash TEXT,
    google_id TEXT UNIQUE,
    google_email TEXT,
    is_owner BOOLEAN NOT NULL DEFAULT FALSE,
    ai_provider TEXT NOT NULL DEFAULT 'deepseek',
    ai_api_keys TEXT NOT NULL DEFAULT '{}',
    sheets_sync_enabled BOOLEAN NOT NULL DEFAULT FALSE,
    house_calls_used INTEGER NOT NULL DEFAULT 0,
    ui_language TEXT NOT NULL DEFAULT 'en',
    explain_language TEXT NOT NULL DEFAULT 'auto',
    ai_relay_base_url TEXT NOT NULL DEFAULT '',
    ai_relay_model TEXT NOT NULL DEFAULT '',
    immersion_target_language TEXT NOT NULL DEFAULT 'follow',
    immersion_source_priority TEXT NOT NULL DEFAULT 'vocab',
    immersion_exclude_proper_nouns BOOLEAN NOT NULL DEFAULT TRUE,
    created_at TIMESTAMPTZ NOT NULL DEFAULT now()
);
ALTER TABLE users ADD COLUMN IF NOT EXISTS house_calls_used INTEGER NOT NULL DEFAULT 0;
ALTER TABLE users ADD COLUMN IF NOT EXISTS ui_language TEXT NOT NULL DEFAULT 'en';
ALTER TABLE users ADD COLUMN IF NOT EXISTS explain_language TEXT NOT NULL DEFAULT 'auto';
ALTER TABLE users ADD COLUMN IF NOT EXISTS ai_relay_base_url TEXT NOT NULL DEFAULT '';
ALTER TABLE users ADD COLUMN IF NOT EXISTS ai_relay_model TEXT NOT NULL DEFAULT '';
ALTER TABLE users ADD COLUMN IF NOT EXISTS immersion_target_language TEXT NOT NULL DEFAULT 'follow';
ALTER TABLE users ADD COLUMN IF NOT EXISTS immersion_source_priority TEXT NOT NULL DEFAULT 'vocab';
ALTER TABLE users ADD COLUMN IF NOT EXISTS immersion_exclude_proper_nouns BOOLEAN NOT NULL DEFAULT TRUE;
CREATE TABLE IF NOT EXISTS house_usage (
    month TEXT PRIMARY KEY,
    calls INTEGER NOT NULL DEFAULT 0,
    cost_usd DOUBLE PRECISION NOT NULL DEFAULT 0
);
CREATE TABLE IF NOT EXISTS sessions (
    token TEXT PRIMARY KEY,
    user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
    expires_at TIMESTAMPTZ NOT NULL
);
CREATE TABLE IF NOT EXISTS documents (
    id TEXT PRIMARY KEY,
    user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    filename TEXT,
    type TEXT,
    content TEXT,
    source_url TEXT,
    learning_language TEXT NOT NULL DEFAULT 'en',
    uploaded_at TIMESTAMPTZ NOT NULL DEFAULT now()
);
ALTER TABLE documents ADD COLUMN IF NOT EXISTS learning_language TEXT NOT NULL DEFAULT 'en';
CREATE TABLE IF NOT EXISTS vocab (
    id TEXT PRIMARY KEY,
    user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    word TEXT,
    sentence TEXT,
    chinese_meaning TEXT,
    ipa TEXT,
    pos TEXT,
    source_doc TEXT,
    date TEXT,
    learning_language TEXT NOT NULL DEFAULT 'en',
    explain_language TEXT NOT NULL DEFAULT 'zh',
    added_at TIMESTAMPTZ NOT NULL DEFAULT now(),
    srs_level INTEGER NOT NULL DEFAULT 0,
    next_review_at TIMESTAMPTZ,
    last_reviewed_at TIMESTAMPTZ
);
ALTER TABLE vocab ADD COLUMN IF NOT EXISTS learning_language TEXT NOT NULL DEFAULT 'en';
ALTER TABLE vocab ADD COLUMN IF NOT EXISTS explain_language TEXT NOT NULL DEFAULT 'zh';
ALTER TABLE vocab ADD COLUMN IF NOT EXISTS srs_level INTEGER NOT NULL DEFAULT 0;
ALTER TABLE vocab ADD COLUMN IF NOT EXISTS next_review_at TIMESTAMPTZ;
ALTER TABLE vocab ADD COLUMN IF NOT EXISTS last_reviewed_at TIMESTAMPTZ;
CREATE TABLE IF NOT EXISTS sentence_notes (
    id TEXT PRIMARY KEY,
    user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    sentence TEXT,
    analysis TEXT,
    source_doc TEXT,
    date TEXT,
    learning_language TEXT NOT NULL DEFAULT 'en',
    explain_language TEXT NOT NULL DEFAULT 'zh',
    added_at TIMESTAMPTZ NOT NULL DEFAULT now()
);
ALTER TABLE sentence_notes ADD COLUMN IF NOT EXISTS learning_language TEXT NOT NULL DEFAULT 'en';
ALTER TABLE sentence_notes ADD COLUMN IF NOT EXISTS explain_language TEXT NOT NULL DEFAULT 'zh';
CREATE TABLE IF NOT EXISTS api_usage (
    user_id TEXT NOT NULL REFERENCES users(id) ON DELETE CASCADE,
    month TEXT NOT NULL,
    calls INTEGER NOT NULL DEFAULT 0,
    cost_usd DOUBLE PRECISION NOT NULL DEFAULT 0,
    PRIMARY KEY (user_id, month)
);
"""


async def get_pool() -> asyncpg.Pool:
    global _pool
    if _pool is None:
        if not DATABASE_URL:
            raise RuntimeError("没有配置 DATABASE_URL，数据库连接不上")
        # statement_cache_size=0：Supabase 的连接池(pgbouncer)在部分模式下不支持
        # 服务端 prepared statement 缓存，关掉这个更保险，不容易踩坑。
        _pool = await asyncpg.create_pool(DATABASE_URL, min_size=1, max_size=5, statement_cache_size=0)
    return _pool


def _serialize_row(row) -> dict:
    d = dict(row)
    for k, v in d.items():
        if isinstance(v, datetime):
            d[k] = v.isoformat()
    return d


def _row_to_user(row) -> Optional[dict]:
    if row is None:
        return None
    d = _serialize_row(row)
    try:
        raw_keys = json.loads(d.get("ai_api_keys") or "{}")
    except json.JSONDecodeError:
        raw_keys = {}
    # AI key 在数据库里是加密存的，读出来这一层统一解密成明文，业务代码不用关心加密细节。
    d["ai_api_keys"] = {k: _decrypt_key(v) for k, v in raw_keys.items()}
    return d


def _parse_dt(s) -> datetime:
    if not s:
        return datetime.now(timezone.utc)
    try:
        return datetime.fromisoformat(s)
    except ValueError:
        return datetime.now(timezone.utc)


async def db_get_user_by_id(user_id: str) -> Optional[dict]:
    pool = await get_pool()
    return _row_to_user(await pool.fetchrow("SELECT * FROM users WHERE id=$1", user_id))


async def db_get_user_by_username(username: str) -> Optional[dict]:
    pool = await get_pool()
    return _row_to_user(await pool.fetchrow("SELECT * FROM users WHERE username=$1", username))


async def db_get_user_by_google_id(google_id: str) -> Optional[dict]:
    pool = await get_pool()
    return _row_to_user(await pool.fetchrow("SELECT * FROM users WHERE google_id=$1", google_id))


async def db_count_users() -> int:
    pool = await get_pool()
    return await pool.fetchval("SELECT count(*) FROM users")


async def db_create_user(
    id: str,
    name: str = "",
    username: Optional[str] = None,
    password_salt: Optional[str] = None,
    password_hash: Optional[str] = None,
    google_id: Optional[str] = None,
    google_email: Optional[str] = None,
    is_owner: bool = False,
    ai_provider: str = "deepseek",
    ai_api_keys: Optional[dict] = None,
    sheets_sync_enabled: bool = False,
    ui_language: str = "en",
) -> dict:
    pool = await get_pool()
    await pool.execute(
        """INSERT INTO users (id, name, username, password_salt, password_hash, google_id, google_email,
                               is_owner, ai_provider, ai_api_keys, sheets_sync_enabled, ui_language)
           VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12)""",
        id, name, username, password_salt, password_hash, google_id, google_email,
        is_owner, ai_provider, json.dumps({k: _encrypt_key(v) for k, v in (ai_api_keys or {}).items()}),
        sheets_sync_enabled, ui_language,
    )
    return await db_get_user_by_id(id)


async def db_update_user_fields(user_id: str, **fields):
    if not fields:
        return
    pool = await get_pool()
    set_parts = []
    values = []
    i = 1
    for k, v in fields.items():
        if k == "ai_api_keys":
            v = json.dumps({pk: _encrypt_key(pv) for pk, pv in v.items()})
        set_parts.append(f"{k}=${i}")
        values.append(v)
        i += 1
    values.append(user_id)
    await pool.execute(f"UPDATE users SET {', '.join(set_parts)} WHERE id=${i}", *values)


SESSION_TTL_DAYS = 30


async def db_create_session(user_id: str, ttl_days: int = SESSION_TTL_DAYS) -> str:
    pool = await get_pool()
    token = secrets.token_urlsafe(32)
    expires_at = datetime.now(timezone.utc) + timedelta(days=ttl_days)
    # 顺手清掉已经过期的旧 session，不用单独搞个定时任务——这张表本来也没多少行。
    await pool.execute("DELETE FROM sessions WHERE expires_at < now()")
    await pool.execute(
        "INSERT INTO sessions (token, user_id, expires_at) VALUES ($1,$2,$3)", token, user_id, expires_at,
    )
    return token


async def db_delete_session(token: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM sessions WHERE token=$1", token)


async def db_get_session_user(token: str) -> Optional[dict]:
    pool = await get_pool()
    row = await pool.fetchrow(
        """SELECT u.* FROM sessions s JOIN users u ON u.id = s.user_id
           WHERE s.token=$1 AND s.expires_at > now()""",
        token,
    )
    return _row_to_user(row)


async def db_create_document(record: dict):
    pool = await get_pool()
    await pool.execute(
        """INSERT INTO documents (id, user_id, filename, type, content, source_url, learning_language, uploaded_at)
           VALUES ($1,$2,$3,$4,$5,$6,$7,$8)""",
        record["id"], record["user_id"], record.get("filename"), record.get("type"),
        record.get("content"), record.get("source_url"), record.get("learning_language", "en"),
        _parse_dt(record.get("uploaded_at")),
    )


async def db_list_documents(user_id: str) -> list:
    pool = await get_pool()
    rows = await pool.fetch("SELECT * FROM documents WHERE user_id=$1 ORDER BY uploaded_at", user_id)
    return [_serialize_row(r) for r in rows]


async def db_delete_documents_for_user(user_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM documents WHERE user_id=$1", user_id)


async def db_get_document(doc_id: str) -> Optional[dict]:
    pool = await get_pool()
    row = await pool.fetchrow("SELECT * FROM documents WHERE id=$1", doc_id)
    return _serialize_row(row) if row else None


async def db_delete_document(doc_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM documents WHERE id=$1", doc_id)


async def db_create_vocab(record: dict):
    pool = await get_pool()
    await pool.execute(
        """INSERT INTO vocab (id, user_id, word, sentence, chinese_meaning, ipa, pos, source_doc, date,
                               learning_language, explain_language, added_at)
           VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12)""",
        record["id"], record["user_id"], record.get("word"), record.get("sentence"),
        record.get("chinese_meaning"), record.get("ipa"), record.get("pos"),
        record.get("source_doc"), record.get("date"),
        record.get("learning_language", "en"), record.get("explain_language", "zh"),
        _parse_dt(record.get("added_at")),
    )


async def db_find_vocab_by_word(user_id: str, word: str) -> Optional[dict]:
    pool = await get_pool()
    row = await pool.fetchrow(
        "SELECT * FROM vocab WHERE user_id=$1 AND lower(word)=lower($2)", user_id, word.strip(),
    )
    return _serialize_row(row) if row else None


async def db_list_vocab(user_id: str) -> list:
    pool = await get_pool()
    rows = await pool.fetch("SELECT * FROM vocab WHERE user_id=$1 ORDER BY added_at", user_id)
    return [_serialize_row(r) for r in rows]


async def db_get_vocab(item_id: str) -> Optional[dict]:
    pool = await get_pool()
    row = await pool.fetchrow("SELECT * FROM vocab WHERE id=$1", item_id)
    return _serialize_row(row) if row else None


async def db_delete_vocab(item_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM vocab WHERE id=$1", item_id)


async def db_delete_vocab_for_user(user_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM vocab WHERE user_id=$1", user_id)


async def db_update_vocab_fields(item_id: str, **fields):
    """照抄 db_update_user_fields 的动态 SET 拼装逻辑，改成按 id 更新 vocab 表。"""
    if not fields:
        return
    pool = await get_pool()
    set_parts = []
    values = []
    i = 1
    for k, v in fields.items():
        set_parts.append(f"{k}=${i}")
        values.append(v)
        i += 1
    values.append(item_id)
    await pool.execute(f"UPDATE vocab SET {', '.join(set_parts)} WHERE id=${i}", *values)


# ---------- 生词复习(简化版莱特纳盒子，不做完整 SM-2) ----------

LEITNER_INTERVALS_DAYS = [1, 3, 7, 14, 30, 90]  # 下标即等级 0-5，等级越高间隔越长


def apply_leitner_result(current_level: int, result: str) -> tuple:
    """返回 (新等级, 下次复习时间)。"认识"升一级(封顶 5)，"不认识"直接归零——
    两种情况都用同一个区间表算下次复习时间，等级 0 对应 1 天，归零后自然明天就会再出现。"""
    new_level = min(current_level + 1, 5) if result == "know" else 0
    next_review_at = datetime.now(timezone.utc) + timedelta(days=LEITNER_INTERVALS_DAYS[new_level])
    return new_level, next_review_at


async def db_count_due_vocab_by_language(user_id: str) -> dict:
    pool = await get_pool()
    rows = await pool.fetch(
        """SELECT learning_language, count(*) AS n FROM vocab
           WHERE user_id=$1 AND (next_review_at IS NULL OR next_review_at <= now())
           GROUP BY learning_language""",
        user_id,
    )
    return {r["learning_language"]: r["n"] for r in rows}


async def db_list_due_vocab(user_id: str, learning_language: str, limit: int) -> list:
    pool = await get_pool()
    rows = await pool.fetch(
        """SELECT * FROM vocab WHERE user_id=$1 AND learning_language=$2
           AND (next_review_at IS NULL OR next_review_at <= now())
           ORDER BY next_review_at NULLS FIRST LIMIT $3""",
        user_id, learning_language, limit,
    )
    return [_serialize_row(r) for r in rows]


async def db_create_sentence_note(record: dict):
    pool = await get_pool()
    await pool.execute(
        """INSERT INTO sentence_notes (id, user_id, sentence, analysis, source_doc, date,
                                        learning_language, explain_language, added_at)
           VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9)""",
        record["id"], record["user_id"], record.get("sentence"), record.get("analysis"),
        record.get("source_doc"), record.get("date"),
        record.get("learning_language", "en"), record.get("explain_language", "zh"),
        _parse_dt(record.get("added_at")),
    )


async def db_list_sentence_notes(user_id: str) -> list:
    pool = await get_pool()
    rows = await pool.fetch("SELECT * FROM sentence_notes WHERE user_id=$1 ORDER BY added_at", user_id)
    return [_serialize_row(r) for r in rows]


async def db_get_sentence_note(item_id: str) -> Optional[dict]:
    pool = await get_pool()
    row = await pool.fetchrow("SELECT * FROM sentence_notes WHERE id=$1", item_id)
    return _serialize_row(row) if row else None


async def db_delete_sentence_note(item_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM sentence_notes WHERE id=$1", item_id)


async def db_delete_sentence_notes_for_user(user_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM sentence_notes WHERE user_id=$1", user_id)


async def db_get_usage(user_id: str, month: str) -> dict:
    pool = await get_pool()
    row = await pool.fetchrow("SELECT calls, cost_usd FROM api_usage WHERE user_id=$1 AND month=$2", user_id, month)
    if not row:
        return {"calls": 0, "cost_usd": 0.0}
    return {"calls": row["calls"], "cost_usd": row["cost_usd"]}


async def db_record_api_call(user_id: str, month: str, cost: float):
    pool = await get_pool()
    await pool.execute(
        """INSERT INTO api_usage (user_id, month, calls, cost_usd) VALUES ($1,$2,1,$3)
           ON CONFLICT (user_id, month) DO UPDATE
           SET calls = api_usage.calls + 1, cost_usd = api_usage.cost_usd + EXCLUDED.cost_usd""",
        user_id, month, cost,
    )


async def db_get_house_usage(month: str) -> dict:
    pool = await get_pool()
    row = await pool.fetchrow("SELECT calls, cost_usd FROM house_usage WHERE month=$1", month)
    if not row:
        return {"calls": 0, "cost_usd": 0.0}
    return {"calls": row["calls"], "cost_usd": row["cost_usd"]}


async def db_record_house_usage(month: str, cost: float):
    pool = await get_pool()
    await pool.execute(
        """INSERT INTO house_usage (month, calls, cost_usd) VALUES ($1,1,$2)
           ON CONFLICT (month) DO UPDATE
           SET calls = house_usage.calls + 1, cost_usd = house_usage.cost_usd + EXCLUDED.cost_usd""",
        month, cost,
    )


async def db_increment_house_calls_used(user_id: str):
    pool = await get_pool()
    await pool.execute("UPDATE users SET house_calls_used = house_calls_used + 1 WHERE id=$1", user_id)


async def db_delete_user(user_id: str):
    pool = await get_pool()
    await pool.execute("DELETE FROM users WHERE id=$1", user_id)


async def db_get_another_user_ordered(exclude_id: str) -> Optional[dict]:
    pool = await get_pool()
    row = await pool.fetchrow(
        "SELECT * FROM users WHERE id != $1 ORDER BY created_at ASC LIMIT 1", exclude_id,
    )
    return _row_to_user(row)


def _normalize_month_usage(month_usage):
    # 兼容改版前的旧格式(以前 month_usage 直接是一个调用次数的数字)——只在从 JSON 迁移时用得到。
    if isinstance(month_usage, (int, float)):
        return {"calls": month_usage, "cost_usd": 0.0}
    return month_usage


async def _migrate_json_to_postgres(conn):
    existing = await conn.fetchval("SELECT count(*) FROM users")
    if existing and existing > 0:
        print("[db] users 表已有数据，跳过 JSON -> Postgres 迁移", flush=True)
        return

    users = _read_json(USERS_FILE, [])
    if not users:
        print("[db] 没有旧的 users.json 数据，数据库从空开始", flush=True)
        return

    print(f"[db] 开始把 {len(users)} 个账号从 JSON 迁移到 Postgres...", flush=True)
    for u in users:
        await conn.execute(
            """INSERT INTO users (id, name, username, password_salt, password_hash, google_id, google_email,
                                   is_owner, ai_provider, ai_api_keys, sheets_sync_enabled)
               VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11)
               ON CONFLICT (id) DO NOTHING""",
            u["id"], u.get("name", ""), u.get("username"), u.get("password_salt"), u.get("password_hash"),
            u.get("google_id"), u.get("google_email"), bool(u.get("is_owner", False)),
            u.get("ai_provider", "deepseek"), json.dumps({k: _encrypt_key(v) for k, v in (u.get("ai_api_keys") or {}).items()}),
            bool(u.get("sheets_sync_enabled", False)),
        )

    documents = _read_json(DOCUMENTS_FILE, [])
    for d in documents:
        await conn.execute(
            """INSERT INTO documents (id, user_id, filename, type, content, source_url, uploaded_at)
               VALUES ($1,$2,$3,$4,$5,$6,$7) ON CONFLICT (id) DO NOTHING""",
            d["id"], d["user_id"], d.get("filename"), d.get("type"), d.get("content"),
            d.get("source_url"), _parse_dt(d.get("uploaded_at")),
        )

    vocab = _read_json(VOCAB_FILE, [])
    for v in vocab:
        await conn.execute(
            """INSERT INTO vocab (id, user_id, word, sentence, chinese_meaning, ipa, pos, source_doc, date, added_at)
               VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10) ON CONFLICT (id) DO NOTHING""",
            v["id"], v["user_id"], v.get("word"), v.get("sentence"), v.get("chinese_meaning"),
            v.get("ipa"), v.get("pos"), v.get("source_doc"), v.get("date"), _parse_dt(v.get("added_at")),
        )

    notes = _read_json(SENTENCE_NOTES_FILE, [])
    for n in notes:
        await conn.execute(
            """INSERT INTO sentence_notes (id, user_id, sentence, analysis, source_doc, date, added_at)
               VALUES ($1,$2,$3,$4,$5,$6,$7) ON CONFLICT (id) DO NOTHING""",
            n["id"], n["user_id"], n.get("sentence"), n.get("analysis"), n.get("source_doc"),
            n.get("date"), _parse_dt(n.get("added_at")),
        )

    usage = _read_json(USAGE_FILE, {})
    for user_id, months in usage.items():
        for month, month_usage in months.items():
            mu = _normalize_month_usage(month_usage)
            await conn.execute(
                """INSERT INTO api_usage (user_id, month, calls, cost_usd) VALUES ($1,$2,$3,$4)
                   ON CONFLICT (user_id, month) DO NOTHING""",
                user_id, month, int(mu.get("calls", 0)), float(mu.get("cost_usd", 0.0)),
            )

    print(
        f"[db] 迁移完成：{len(users)} 账号, {len(documents)} 文档, {len(vocab)} 生词, {len(notes)} 句子笔记",
        flush=True,
    )

    # 迁移完成后把旧 JSON 文件改名备份，不删除——留个痕迹，也避免下次启动误判。
    for f in [USERS_FILE, DOCUMENTS_FILE, VOCAB_FILE, SENTENCE_NOTES_FILE, USAGE_FILE, SESSIONS_FILE]:
        if f.exists():
            f.rename(f.with_suffix(f.suffix + ".migrated"))


async def _migrate_encrypt_existing_keys(conn):
    # 一次性/幂等的补丁：把 users.ai_api_keys 里还没加密过的明文 key 加密。已经加密过的
    # 值直接能被 _fernet.decrypt 解出来，会跳过，所以这个函数每次启动跑一遍也没问题。
    rows = await conn.fetch("SELECT id, ai_api_keys FROM users")
    changed = 0
    for row in rows:
        try:
            keys = json.loads(row["ai_api_keys"] or "{}")
        except json.JSONDecodeError:
            continue
        new_keys = {}
        row_changed = False
        for k, v in keys.items():
            if not v or _is_encrypted(v):
                new_keys[k] = v
            else:
                new_keys[k] = _encrypt_key(v)
                row_changed = True
        if row_changed:
            await conn.execute("UPDATE users SET ai_api_keys=$1 WHERE id=$2", json.dumps(new_keys), row["id"])
            changed += 1
    if changed:
        print(f"[db] 把 {changed} 个账号的 AI key 从明文改成加密存储", flush=True)


async def init_db():
    _migrate_legacy_invite_users()
    pool = await get_pool()
    async with pool.acquire() as conn:
        await conn.execute(SCHEMA_SQL)
        await _migrate_json_to_postgres(conn)
        await _migrate_encrypt_existing_keys(conn)


@app.on_event("startup")
async def on_startup():
    await init_db()


# ---------- 用户 / 账号注册登录(用户名密码 + Google 登录，登录后发一个 session token) ----------
# 谁都能自己注册用户名 + 密码，用户名/密码没有任何格式限制。第一个注册/登录的账号自动成为
# 主账号(唯一能开 Google Sheets 同步的账号)。登录成功后发一个 token，之后每次请求带
# Authorization: Bearer <token>；token 存在 Postgres 的 sessions 表里，退出登录时会把
# 对应的 token 删掉。

bearer_scheme = HTTPBearer(auto_error=False)


def _hash_password(password: str, salt: str) -> str:
    return hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 100_000).hex()


@app.get("/api/config")
async def get_public_config():
    # Turnstile site key 本来就是要贴到前端页面里的公开值，不是敏感信息，
    # 走接口发而不是写死在 HTML 里，只是为了配置还是统一走环境变量。
    return {"turnstile_site_key": TURNSTILE_SITE_KEY}


async def _verify_turnstile(token: str, remote_ip: str = "") -> bool:
    if not TURNSTILE_SECRET_KEY:
        return True  # 没配置密钥就跳过校验(本地开发环境没必要强制装验证码)
    if not token:
        return False
    async with httpx.AsyncClient(timeout=10) as client:
        resp = await client.post(
            "https://challenges.cloudflare.com/turnstile/v0/siteverify",
            data={"secret": TURNSTILE_SECRET_KEY, "response": token, "remoteip": remote_ip},
        )
    if resp.status_code != 200:
        return False
    return bool(resp.json().get("success"))


MIN_PASSWORD_LENGTH = 8


def _validate_password_strength(password: str):
    if len(password) < MIN_PASSWORD_LENGTH:
        raise HTTPException(400, f"密码至少要 {MIN_PASSWORD_LENGTH} 位")


def _migrate_legacy_invite_users():
    """把改版前"邀请码即凭证"的老账号，自动补上 username/password(初始密码沿用原邀请码)。
    这个函数只操作 JSON 文件，在 Postgres 迁移之前跑，保证迁移过去的数据里没有这种老格式。"""
    users = _read_json(USERS_FILE, [])
    if not users:
        return
    print(
        "[startup] users.json 现有 "
        + str(len(users))
        + " 个账号: "
        + ", ".join(
            f"{u.get('username') or u.get('name')}(密码={'有' if u.get('password_hash') else '无'})" for u in users
        ),
        flush=True,
    )
    existing_usernames = {u.get("username") for u in users if u.get("username")}
    changed = False
    for u in users:
        if u.get("password_hash") or not u.get("invite_code"):
            continue
        base = re.sub(r"[^a-zA-Z0-9]", "", u.get("name", "")).lower() or "user"
        username = base
        n = 1
        while username in existing_usernames:
            n += 1
            username = f"{base}{n}"
        existing_usernames.add(username)
        salt = secrets.token_hex(16)
        u["username"] = username
        u["password_salt"] = salt
        u["password_hash"] = _hash_password(u["invite_code"], salt)
        changed = True
        print(f"[migrate] 账号「{u.get('name','')}」自动分配用户名：{username}（初始密码是原来的邀请码）", flush=True)
    if changed:
        _write_json(USERS_FILE, users)


def _audit(event: str, **fields):
    # 排查"密码莫名其妙变了"这类问题用的审计日志，打到 Railway 的部署日志里，
    # 不落盘、不影响正常功能——纯粹是为了下次万一再出问题时有据可查。
    parts = " ".join(f"{k}={v}" for k, v in fields.items())
    print(f"[audit] {datetime.now(timezone.utc).isoformat()} {event} {parts}", flush=True)


class RegisterRequest(BaseModel):
    username: str
    password: str
    turnstile_token: str = ""
    ui_language: str = "en"


@app.post("/api/register")
@limiter.limit("10/minute")
async def register(request: Request, req: RegisterRequest):
    username = req.username.strip()
    password = req.password
    if not username or not password:
        raise HTTPException(400, "用户名和密码不能为空")
    if not await _verify_turnstile(req.turnstile_token, get_remote_address(request)):
        raise HTTPException(400, "人机验证没通过，刷新页面重试一下")
    _validate_password_strength(password)

    if await db_get_user_by_username(username):
        raise HTTPException(400, "这个用户名已经被注册了，换一个")

    salt = secrets.token_hex(16)
    user_count = await db_count_users()
    new_user = await db_create_user(
        id="u_" + uuid.uuid4().hex[:10],
        name=username,
        username=username,
        password_salt=salt,
        password_hash=_hash_password(password, salt),
        is_owner=(user_count == 0),
        ui_language=req.ui_language if req.ui_language in UI_LANGUAGES else "en",
    )
    _audit("register", user_id=new_user["id"], username=username)
    token = await db_create_session(new_user["id"])
    return {
        "token": token,
        "id": new_user["id"],
        "name": username,
        "is_owner": new_user["is_owner"],
        "ui_language": new_user.get("ui_language", "en"),
        "house_trial_enabled": bool(HOUSE_AI_API_KEY),
        "house_calls_total": HOUSE_FREE_CALLS_PER_USER,
    }


class LoginRequest(BaseModel):
    username: str
    password: str
    turnstile_token: str = ""


@app.post("/api/login")
@limiter.limit("10/minute")
async def login(request: Request, req: LoginRequest):
    if not await _verify_turnstile(req.turnstile_token, get_remote_address(request)):
        raise HTTPException(400, "人机验证没通过，刷新页面重试一下")
    user = await db_get_user_by_username(req.username.strip())
    if not user or not user.get("password_hash"):
        raise HTTPException(401, "用户名或密码不对")
    expected = _hash_password(req.password, user.get("password_salt") or "")
    if not hmac.compare_digest(expected, user["password_hash"]):
        raise HTTPException(401, "用户名或密码不对")
    token = await db_create_session(user["id"])
    return {
        "token": token,
        "id": user["id"],
        "name": user.get("name", ""),
        "is_owner": user.get("is_owner", False),
        "ui_language": user.get("ui_language", "en"),
    }


@app.post("/api/logout")
async def logout(credentials: Optional[HTTPAuthorizationCredentials] = Depends(bearer_scheme)):
    if credentials:
        await db_delete_session(credentials.credentials)
    return {"ok": True}


async def get_current_user(credentials: Optional[HTTPAuthorizationCredentials] = Depends(bearer_scheme)):
    if not credentials:
        raise HTTPException(401, "未登录")
    user = await db_get_session_user(credentials.credentials)
    if not user:
        raise HTTPException(401, "登录状态已失效，重新登录一下")
    return user


@app.get("/api/me")
async def get_me(user: dict = Depends(get_current_user)):
    return {
        "id": user["id"],
        "name": user.get("name", ""),
        "is_owner": user.get("is_owner", False),
        "ui_language": user.get("ui_language", "en"),
    }


class ChangePasswordRequest(BaseModel):
    new_password: str


@app.post("/api/change-password")
@limiter.limit("10/minute")
async def change_password(request: Request, req: ChangePasswordRequest, user: dict = Depends(get_current_user)):
    _validate_password_strength(req.new_password)
    salt = secrets.token_hex(16)
    await db_update_user_fields(user["id"], password_salt=salt, password_hash=_hash_password(req.new_password, salt))
    _audit("change_password", user_id=user["id"], username=user.get("username"))
    return {"ok": True}


class ChangeUsernameRequest(BaseModel):
    new_username: str


@app.post("/api/change-username")
async def change_username(req: ChangeUsernameRequest, user: dict = Depends(get_current_user)):
    new_username = req.new_username.strip()
    if not new_username:
        raise HTTPException(400, "新用户名不能为空")
    existing = await db_get_user_by_username(new_username)
    if existing and existing["id"] != user["id"]:
        raise HTTPException(400, "这个用户名已经被占用了，换一个")
    await db_update_user_fields(user["id"], username=new_username, name=new_username)
    _audit("change_username", user_id=user["id"], old_username=user.get("username"), new_username=new_username)
    return {"ok": True, "username": new_username, "name": new_username}


@app.get("/api/account/export")
@limiter.limit("5/minute")
async def export_account_data(request: Request, user: dict = Depends(get_current_user)):
    # 导出的是"你自己能看到的数据"，不包含密码哈希、AI key 这类凭证信息——没必要导出，
    # 导出反而多一份泄露风险。
    documents = await db_list_documents(user["id"])
    vocab = await db_list_vocab(user["id"])
    notes = await db_list_sentence_notes(user["id"])
    month_key = datetime.now().strftime("%Y-%m")
    usage = await db_get_usage(user["id"], month_key)
    return {
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "account": {
            "id": user["id"],
            "username": user.get("username"),
            "name": user.get("name"),
            "google_email": user.get("google_email") or None,
            "is_owner": user.get("is_owner", False),
            "ai_provider": user.get("ai_provider"),
            "sheets_sync_enabled": user.get("sheets_sync_enabled", False),
        },
        "documents": documents,
        "vocab": vocab,
        "sentence_notes": notes,
        "usage_this_month": usage,
    }


@app.delete("/api/account")
@limiter.limit("5/minute")
async def delete_account(request: Request, user: dict = Depends(get_current_user)):
    # 主账号删号的话先把主账号身份转给另一个还在的账号(按注册时间挑最早的那个)，
    # 不然 Google Sheets 同步开关这些主账号专属功能就没人能用了。
    if user.get("is_owner"):
        successor = await db_get_another_user_ordered(user["id"])
        if successor:
            await db_update_user_fields(successor["id"], is_owner=True)
            _audit("owner_transfer", from_user_id=user["id"], to_user_id=successor["id"])
    await db_delete_user(user["id"])
    _audit("account_deleted", user_id=user["id"], username=user.get("username"))
    return {"ok": True}


# ---------- Google 登录 / 关联 ----------
# 谷歌登录默认按 google_id 查/建账号，跟用户名密码账号是分开的。要让"同一个账号既能用
# 用户名密码登录、也能用 Google 登录"，需要显式关联：已登录用户在设置面板点"关联 Google
# 账号"，走一遍 OAuth，把拿到的 google_id 写到当前账号上，而不是新建/登录到别的账号。
# 由于 OAuth 是整页跳转，中途拿不到 Authorization header，所以先用一个短期一次性 nonce
# (通过已登录状态的 POST 请求换到)把"是谁在发起关联"这件事传过去，再带着 nonce 跳转。

_pending_oauth_states: dict = {}  # state -> {"ts": 时间, "link_user_id": 关联目标账号id或None}
_pending_link_nonces: dict = {}  # nonce -> {"ts": 时间, "user_id": ...}


def _cleanup_stale(d: dict, max_age_seconds: int = 600):
    now = datetime.now(timezone.utc)
    stale = [k for k, v in d.items() if (now - v["ts"]).total_seconds() > max_age_seconds]
    for k in stale:
        d.pop(k, None)


def _new_oauth_state(link_user_id: Optional[str] = None) -> str:
    _cleanup_stale(_pending_oauth_states)
    state = secrets.token_urlsafe(16)
    _pending_oauth_states[state] = {"ts": datetime.now(timezone.utc), "link_user_id": link_user_id}
    return state


@app.post("/api/auth/google/link-init")
@limiter.limit("10/minute")
async def google_link_init(request: Request, user: dict = Depends(get_current_user)):
    _cleanup_stale(_pending_link_nonces)
    nonce = secrets.token_urlsafe(24)
    _pending_link_nonces[nonce] = {"ts": datetime.now(timezone.utc), "user_id": user["id"]}
    return {"nonce": nonce}


@app.get("/api/auth/google/login")
async def google_login(request: Request, link_nonce: str = ""):
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(500, "这个部署还没配置 Google 登录(缺 GOOGLE_CLIENT_ID)")

    link_user_id = None
    if link_nonce:
        entry = _pending_link_nonces.pop(link_nonce, None)
        if not entry:
            raise HTTPException(400, "关联请求已过期，回设置里重新点一次")
        link_user_id = entry["user_id"]

    state = _new_oauth_state(link_user_id=link_user_id)
    redirect_uri = str(request.base_url) + "api/auth/google/callback"
    params = {
        "client_id": GOOGLE_CLIENT_ID,
        "redirect_uri": redirect_uri,
        "response_type": "code",
        "scope": "openid email profile",
        "state": state,
        "prompt": "select_account",
    }
    return RedirectResponse("https://accounts.google.com/o/oauth2/v2/auth?" + urlencode(params))


@app.get("/api/auth/google/callback")
async def google_callback(request: Request, code: str = "", state: str = "", error: str = ""):
    if error:
        raise HTTPException(400, f"Google 登录失败: {error}")
    if not state or state not in _pending_oauth_states:
        raise HTTPException(400, "登录状态已过期，重新点一次「用 Google 登录」")
    state_entry = _pending_oauth_states.pop(state)
    link_user_id = state_entry.get("link_user_id")

    redirect_uri = str(request.base_url) + "api/auth/google/callback"
    async with httpx.AsyncClient(timeout=15) as client:
        token_resp = await client.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            },
        )
    if token_resp.status_code != 200:
        raise HTTPException(400, f"Google 登录换取 token 失败: {token_resp.text}")
    google_access_token = token_resp.json().get("access_token", "")

    async with httpx.AsyncClient(timeout=15) as client:
        profile_resp = await client.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {google_access_token}"},
        )
    if profile_resp.status_code != 200:
        raise HTTPException(400, "拿不到 Google 账号信息")
    profile = profile_resp.json()
    google_id = profile.get("sub", "")
    email = profile.get("email", "")
    name = profile.get("name") or email or "Google 用户"
    if not google_id:
        raise HTTPException(400, "Google 账号信息里没有用户 ID")

    if link_user_id:
        # 关联流程：把这个 google_id 绑到当前登录的账号上，而不是新建/登录到别的账号。
        # 如果这个 google_id 之前绑在另一个账号上(比如之前误建了个空的 Google 账号)，
        # 就把它从旧账号上摘下来，改绑到这次要关联的账号——反正 Google 身份认证本身
        # 已经证明了操作者对这个 Google 账号有控制权，重新指向不算越权。
        target = await db_get_user_by_id(link_user_id)
        if not target:
            raise HTTPException(400, "要关联的账号不存在了，重新登录后再试一次")
        existing_owner = await db_get_user_by_google_id(google_id)
        if existing_owner and existing_owner["id"] != link_user_id:
            await db_update_user_fields(existing_owner["id"], google_id=None, google_email=None)
            _audit("google_unlink", user_id=existing_owner["id"], google_email=email)
        await db_update_user_fields(link_user_id, google_id=google_id, google_email=email)
        _audit("google_link", user_id=link_user_id, google_email=email)
        token = await db_create_session(link_user_id)
        return RedirectResponse(f"/#token={token}&google_linked=1")

    user = await db_get_user_by_google_id(google_id)
    if not user:
        user_count = await db_count_users()
        user = await db_create_user(
            id="u_" + uuid.uuid4().hex[:10],
            name=name,
            google_id=google_id,
            google_email=email,
            is_owner=(user_count == 0),
        )
        _audit("google_register", user_id=user["id"], google_email=email)

    token = await db_create_session(user["id"])
    return RedirectResponse(f"/#token={token}")


# ---------- 文档管理(粘贴文本 / 网址导入 / PDF·DOCX 上传，最终都存成统一的文字文档) ----------

# 学习语言自动检测：只按 Unicode 字符区间做粗略判断，不依赖任何第三方语言检测库。
# 对日语/韩语/中文/俄语/阿拉伯语/泰语这些"文字系统本身就不一样"的语言，字符区间足够准确；
# 拉丁字母语系(英/法/西/德/意/葡...)彼此没法靠字符区分，统一先猜成英语，
# 用户在添加文章的时候可以在下拉框里手动改成正确的语言。
_SCRIPT_RANGES = [
    ("ja", re.compile(r"[぀-ヿ]")),  # 平假名/片假名，只要出现就几乎能确定是日语
    ("ko", re.compile(r"[가-힣]")),  # 谚文
    ("zh", re.compile(r"[一-鿿]")),  # CJK 统一表意文字(没有假名的情况下判定为中文)
    ("ru", re.compile(r"[Ѐ-ӿ]")),  # 西里尔字母
    ("ar", re.compile(r"[؀-ۿ]")),  # 阿拉伯字母
    ("th", re.compile(r"[฀-๿]")),  # 泰文
]


def detect_learning_language(text: str) -> str:
    sample = text[:2000]
    for lang, pattern in _SCRIPT_RANGES:
        if pattern.search(sample):
            return lang
    return "en"


# pdfplumber 自带的 page.extract_text() 判断"两个字符间要不要插空格"用的是固定阈值
# (x_tolerance 默认 3pt)，双栏论文常见的 8-10pt 小字号下，正常词间距往往就只有
# 1.5-2pt，比这个固定阈值还小，于是大量单词被直接拼在一起、读不出空格。
# 这里改成按当前字符字号的比例算阈值(实测同一份文档里，同词内相邻字符间距基本是
# 0，词与词之间大约是字号的 0.2 倍左右)，取 0.15 留出安全余量，能可靠区分"词内
# 字符"和"词与词之间"，不受字号大小影响。
_PDF_WORD_GAP_RATIO = 0.15
# 双栏排版另一个问题是 pdfplumber 按整页宽度从左到右、从上到下读字符，同一行高度上
# 左右两栏的文字会被交替拼在一起(比如左栏第一行紧跟着右栏第一行)，读出来的顺序完全
# 是乱的。下面先检测页面中间有没有一条贯穿大半页高、几乎没有文字经过的竖直空白带
# (栏间距)，有的话就分左右两栏分别按顺序提取，没有就当单栏处理。


def _pdf_cluster_lines(chars, y_tolerance=3):
    if not chars:
        return []
    chars = sorted(chars, key=lambda c: (c["top"], c["x0"]))
    lines = []
    current_line = [chars[0]]
    current_top = chars[0]["top"]
    for ch in chars[1:]:
        if abs(ch["top"] - current_top) <= y_tolerance:
            current_line.append(ch)
            current_top = (current_top + ch["top"]) / 2
        else:
            lines.append(sorted(current_line, key=lambda c: c["x0"]))
            current_line = [ch]
            current_top = ch["top"]
    lines.append(sorted(current_line, key=lambda c: c["x0"]))
    return lines


def _pdf_line_to_text(line_chars, gap_ratio=_PDF_WORD_GAP_RATIO):
    if not line_chars:
        return ""
    parts = [line_chars[0]["text"]]
    for prev, cur in zip(line_chars, line_chars[1:]):
        if cur["text"].isspace() or prev["text"].isspace():
            parts.append(cur["text"])
            continue
        gap = cur["x0"] - prev["x1"]
        size = cur.get("size") or prev.get("size") or 10
        if gap > size * gap_ratio:
            parts.append(" ")
        parts.append(cur["text"])
    return "".join(parts)


def _pdf_lines_to_text(chars, y_tolerance=3, gap_ratio=_PDF_WORD_GAP_RATIO):
    lines = _pdf_cluster_lines(chars, y_tolerance=y_tolerance)
    return "\n".join(_pdf_line_to_text(line, gap_ratio=gap_ratio) for line in lines)


def _pdf_find_column_gutter(chars, page_width, search_lo=0.28, search_hi=0.72, bucket=1.0, min_gap_width=6):
    """在页面横向 28%-72% 范围内找一条贯穿的空白带，当作双栏的分栏线。
    范围特意避开页面最左/最右，防止把普通的左右页边距误判成栏间距。"""
    if not chars:
        return None
    n_buckets = int(page_width / bucket) + 1
    covered = [False] * n_buckets

    def mark(x0, x1):
        lo = max(0, int(x0 / bucket))
        hi = min(n_buckets - 1, int(x1 / bucket))
        for i in range(lo, hi + 1):
            covered[i] = True

    for c in chars:
        mark(c["x0"], c["x1"])

    lo_bucket = int(page_width * search_lo / bucket)
    hi_bucket = int(page_width * search_hi / bucket)

    best_start, best_len = None, 0
    run_start = None
    for i in range(lo_bucket, hi_bucket + 1):
        if not covered[i]:
            if run_start is None:
                run_start = i
        else:
            if run_start is not None:
                run_len = i - run_start
                if run_len > best_len:
                    best_len, best_start = run_len, run_start
                run_start = None
    if run_start is not None:
        run_len = hi_bucket + 1 - run_start
        if run_len > best_len:
            best_len, best_start = run_len, run_start

    if best_start is None or best_len * bucket < min_gap_width:
        return None
    return (best_start + best_len / 2) * bucket


def _extract_pdf_page_text(page) -> str:
    chars = page.chars
    if not chars:
        return ""
    gutter = _pdf_find_column_gutter(chars, page.width)
    if gutter is None:
        return _pdf_lines_to_text(chars)
    left = [c for c in chars if c["x1"] <= gutter]
    right = [c for c in chars if c["x0"] >= gutter]
    # 极少数字符可能正好压在分栏线上(比如跨栏的图/表标题)，归到左栏，不丢字
    stray = [c for c in chars if c not in left and c not in right]
    left += stray
    return _pdf_lines_to_text(left) + "\n\n" + _pdf_lines_to_text(right)


def extract_pdf_text(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = [_extract_pdf_page_text(page) for page in pdf.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())


def extract_docx_text(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


@app.post("/api/upload")
async def upload_document(
    file: UploadFile = File(...),
    learning_language: str = Form(""),
    user: dict = Depends(get_current_user),
):
    ext = Path(file.filename).suffix.lower()
    if ext not in (".pdf", ".docx"):
        raise HTTPException(400, "只支持 PDF 或 DOCX 文件")

    file_bytes = await file.read()
    try:
        if ext == ".pdf":
            content = await asyncio.to_thread(extract_pdf_text, file_bytes)
        else:
            content = await asyncio.to_thread(extract_docx_text, file_bytes)
    except Exception as e:
        raise HTTPException(400, f"解析文件失败：{e}")

    if not content.strip():
        raise HTTPException(400, "没能从这个文件里提取出文字，可能是扫描版 PDF(图片形式，没有文字层)")

    doc_id = uuid.uuid4().hex[:12]
    record = {
        "id": doc_id,
        "user_id": user["id"],
        "filename": Path(file.filename).stem,
        "type": "text",
        "content": content,
        "learning_language": learning_language.strip() or detect_learning_language(content),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }
    await db_create_document(record)
    return record


class PasteRequest(BaseModel):
    title: str = ""
    content: str
    learning_language: str = ""


@app.post("/api/paste")
async def paste_document(req: PasteRequest, user: dict = Depends(get_current_user)):
    if not req.content.strip():
        raise HTTPException(400, "文章内容不能为空")

    doc_id = uuid.uuid4().hex[:12]
    record = {
        "id": doc_id,
        "user_id": user["id"],
        "filename": req.title.strip() or f"粘贴文章-{doc_id}",
        "type": "text",
        "content": req.content,
        "learning_language": req.learning_language.strip() or detect_learning_language(req.content),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }
    await db_create_document(record)
    return record


class UrlFetchRequest(BaseModel):
    url: str
    learning_language: str = ""


@app.post("/api/fetch-url")
async def fetch_url_document(req: UrlFetchRequest, user: dict = Depends(get_current_user)):
    url = req.url.strip()
    if not url.startswith("http://") and not url.startswith("https://"):
        raise HTTPException(400, "网址格式不对，需要以 http:// 或 https:// 开头")

    downloaded = await asyncio.to_thread(trafilatura.fetch_url, url)
    if not downloaded:
        raise HTTPException(400, "打不开这个网址，检查一下网址是否正确，或者这个网站限制了访问")

    extracted = await asyncio.to_thread(
        trafilatura.extract,
        downloaded,
        include_comments=False,
        include_tables=False,
        favor_precision=True,
        with_metadata=True,
        output_format="json",
    )
    if not extracted:
        raise HTTPException(400, "抓到了页面，但没能提取出正文，可能这个页面不是文章页")

    data = json.loads(extracted)
    title = (data.get("title") or "").strip() or url
    content = (data.get("text") or "").strip()
    if not content:
        raise HTTPException(400, "没有提取到正文内容")

    doc_id = uuid.uuid4().hex[:12]
    record = {
        "id": doc_id,
        "user_id": user["id"],
        "filename": title,
        "type": "text",
        "content": content,
        "source_url": url,
        "learning_language": req.learning_language.strip() or detect_learning_language(content),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }
    await db_create_document(record)
    return record


@app.get("/api/documents")
async def list_documents(user: dict = Depends(get_current_user)):
    return await db_list_documents(user["id"])


@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str, user: dict = Depends(get_current_user)):
    target = await db_get_document(doc_id)
    if not target or target.get("user_id") != user["id"]:
        raise HTTPException(404, "没找到这篇文章")
    await db_delete_document(doc_id)
    # 生词/句子笔记是按文章标题(source_doc)关联的，不是按这篇文章的 id，
    # 删文章不会跟着删掉已经存的生词——跟本地删生词不会删 Sheet 里对应行是同一个道理，
    # 避免"删错一下就连带丢数据"。
    return {"ok": True}


@app.get("/api/coverage/{doc_id}")
async def get_coverage(doc_id: str, user: dict = Depends(get_current_user)):
    """中/日/韩文章的生词覆盖率——拉丁字母语言前端自己算，不用调这个接口。"""
    doc = await db_get_document(doc_id)
    if not doc or doc.get("user_id") != user["id"]:
        raise HTTPException(404, "没找到这篇文章")
    lang = doc.get("learning_language", "en")
    tokenizer = COVERAGE_TOKENIZERS.get(lang)
    if not tokenizer:
        raise HTTPException(400, f"这个语言不支持覆盖率统计: {lang}")

    vocab_rows = await db_list_vocab(user["id"])
    known_words = {v["word"].strip().lower() for v in vocab_rows if v.get("word") and v.get("learning_language") == lang}

    paragraphs = [p.strip() for p in re.split(r"\n+", doc.get("content") or "") if p.strip()]
    total_tokens = 0
    known_tokens = 0
    for p in paragraphs:
        for token in tokenizer(p):
            total_tokens += 1
            if token.strip().lower() in known_words:
                known_tokens += 1

    coverage_pct = round(known_tokens / total_tokens * 100) if total_tokens else 0
    return {"total_tokens": total_tokens, "known_tokens": known_tokens, "coverage_pct": coverage_pct}


# ---------- 界面语言 / AI 讲解语言 / 学习语言 ----------
# ui_language：界面文案用哪种语言，目前只开放中/英两个选项。
# explain_language：AI 讲解输出用哪种语言，'auto' 表示跟随 ui_language 实时计算，
# 用户也可以手动固定成某一种，不受界面语言变化影响。
# learning_language：某一篇文章/某一条生词笔记正在学习的目标语言，开放式的，不是只有中英两个选项。
UI_LANGUAGES = ["zh", "en", "ko"]
EXPLAIN_LANGUAGE_CHOICES = ["auto", "zh", "en"]

LANGUAGE_LABELS = {
    "zh": "中文", "en": "英文", "ja": "日语", "ko": "韩语", "fr": "法语",
    "es": "西班牙语", "de": "德语", "it": "意大利语", "pt": "葡萄牙语",
    "ru": "俄语", "ar": "阿拉伯语", "th": "泰语",
}


def resolve_explain_language(user: dict) -> str:
    explain = user.get("explain_language", "auto")
    if explain == "auto":
        return user.get("ui_language", "en")
    return explain


# ---------- AI 调用(多服务商：DeepSeek / OpenAI / Claude / Gemini，用各自用户自己填的 key) ----------

# 各家 API 会定期弃用旧模型名(2026 年上半年 deepseek-chat / claude-3.5-haiku / gemini-1.5-flash /
# gpt-4o-mini 都先后失效过)，报 "不支持的模型名 / model not found" 这类错就是这里过期了，
# 更新成各家当前的轻量档模型即可。Gemini 用官方的 -latest 别名自动跟随最新版，不用手动追。
PROVIDER_CONFIG = {
    "deepseek": {"label": "DeepSeek", "kind": "openai", "url": "https://api.deepseek.com/chat/completions", "model": "deepseek-v4-flash"},
    "openai": {"label": "OpenAI", "kind": "openai", "url": "https://api.openai.com/v1/chat/completions", "model": "gpt-5-mini"},
    "claude": {"label": "Claude (Anthropic)", "kind": "claude", "model": "claude-haiku-4-5-20251001"},
    "gemini": {"label": "Gemini (Google)", "kind": "gemini", "model": "gemini-flash-latest"},
}

# 每百万 token 的价格(美元)，按写这段代码时各家官网公布的价格粗略估算，仅供参考——
# 实际计费以服务商账单为准，价格会变，这里不会自动跟着更新。
PROVIDER_PRICING = {
    "deepseek": {"input": 0.14, "output": 0.28},
    "openai": {"input": 0.25, "output": 2.00},
    "claude": {"input": 1.00, "output": 5.00},
    "gemini": {"input": 0.30, "output": 2.50},
}


async def record_api_call(user_id: str, provider: str = "", input_tokens: int = 0, output_tokens: int = 0) -> float:
    month_key = datetime.now().strftime("%Y-%m")
    rates = PROVIDER_PRICING.get(provider)
    cost = 0.0
    if rates:
        cost = (input_tokens / 1_000_000) * rates["input"] + (output_tokens / 1_000_000) * rates["output"]
    await db_record_api_call(user_id, month_key, cost)
    return cost


@app.get("/api/usage")
async def get_usage(user: dict = Depends(get_current_user)):
    month_key = datetime.now().strftime("%Y-%m")
    month_usage = await db_get_usage(user["id"], month_key)
    return {
        "month": month_key,
        "count": month_usage.get("calls", 0),
        "cost_usd": round(month_usage.get("cost_usd", 0.0), 4),
    }


async def _call_openai_compatible(url: str, model: str, api_key: str, prompt: str, json_mode: bool):
    payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.3}
    if json_mode:
        payload["response_format"] = {"type": "json_object"}
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.post(url, headers={"Authorization": f"Bearer {api_key}"}, json=payload)
    if resp.status_code != 200:
        raise HTTPException(502, f"AI 接口调用失败: {resp.text}")
    try:
        data = resp.json()
        usage = data.get("usage", {})
        return data["choices"][0]["message"]["content"], usage.get("prompt_tokens", 0), usage.get("completion_tokens", 0)
    except (json.JSONDecodeError, KeyError, IndexError):
        raise HTTPException(502, f"AI 返回格式异常: {resp.text}")


async def _call_claude(api_key: str, prompt: str, json_mode: bool):
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": PROVIDER_CONFIG["claude"]["model"],
                "max_tokens": 1024,
                "messages": [{"role": "user", "content": prompt}],
            },
        )
    if resp.status_code != 200:
        raise HTTPException(502, f"AI 接口调用失败(Claude): {resp.text}")
    data = resp.json()
    usage = data.get("usage", {})
    try:
        return data["content"][0]["text"], usage.get("input_tokens", 0), usage.get("output_tokens", 0)
    except (KeyError, IndexError):
        raise HTTPException(502, f"Claude 返回格式异常: {data}")


async def _call_gemini(api_key: str, prompt: str, json_mode: bool):
    model = PROVIDER_CONFIG["gemini"]["model"]
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.post(url, json={"contents": [{"parts": [{"text": prompt}]}]})
    if resp.status_code != 200:
        raise HTTPException(502, f"AI 接口调用失败(Gemini): {resp.text}")
    data = resp.json()
    usage = data.get("usageMetadata", {})
    try:
        text = data["candidates"][0]["content"]["parts"][0]["text"]
        return text, usage.get("promptTokenCount", 0), usage.get("candidatesTokenCount", 0)
    except (KeyError, IndexError):
        raise HTTPException(502, f"Gemini 返回格式异常: {data}")


async def call_ai(prompt: str, provider: str, api_key: str, user_id: str, json_mode: bool = False,
                   relay_base_url: str = "", relay_model: str = ""):
    """返回 (回复文本, 这次调用的估算花费 usd)。relay_base_url/relay_model 只在用户自己配置了 key 时
    生效，绝不能用在公共体验 key 上，否则用户能把站长的 key 通过自己指定的中转地址偷偷转发出去。
    配了中转地址后，四个服务商统一走 OpenAI 兼容接口转发——绝大多数"中转站"不管实际代理的
    是哪家模型，对外都是暴露一个统一的 OpenAI 兼容接口，所以这么处理兼容性最好；没配中转
    地址时，Claude/Gemini 仍然各自走官方原生接口(格式跟 OpenAI 完全不同，没法直接对调)。
    relay_model：不同中转站给同一个服务商用的模型名可能跟官方不一样，配了就用这个名字
    覆盖默认值，没配就还是用 PROVIDER_CONFIG 里的官方模型名。"""
    if not api_key:
        raise HTTPException(400, "还没有配置 AI API Key，先去设置里填一下")
    cfg = PROVIDER_CONFIG.get(provider)
    if not cfg:
        raise HTTPException(400, f"不支持的 AI 服务商: {provider}")

    if relay_base_url:
        url = f"{relay_base_url.rstrip('/')}/chat/completions"
        model = relay_model.strip() or cfg["model"]
        text, in_tok, out_tok = await _call_openai_compatible(url, model, api_key, prompt, json_mode)
    elif cfg["kind"] == "openai":
        text, in_tok, out_tok = await _call_openai_compatible(cfg["url"], cfg["model"], api_key, prompt, json_mode)
    elif cfg["kind"] == "claude":
        text, in_tok, out_tok = await _call_claude(api_key, prompt, json_mode)
    elif cfg["kind"] == "gemini":
        text, in_tok, out_tok = await _call_gemini(api_key, prompt, json_mode)
    else:
        raise HTTPException(500, "AI 服务商配置错误")

    cost = await record_api_call(user_id, provider, in_tok, out_tok)
    return text, cost


async def resolve_ai_credentials(user: dict):
    """决定这次调用实际用谁的 key：优先用用户自己配置的；没配的话，如果还有公共
    体验额度(账户没用满 10 次、当月公共预算没超)，就用站长的公共 key。
    返回 (provider, api_key, using_house, blocked_reason)。api_key 为空时，
    blocked_reason 是 "no_key" / "user_limit" / "global_budget" 之一，方便上层给出准确的提示。"""
    provider = user.get("ai_provider", "deepseek")
    own_key = user.get("ai_api_keys", {}).get(provider, "")
    if own_key:
        return provider, own_key, False, None

    if not HOUSE_AI_API_KEY:
        return provider, "", False, "no_key"
    if user.get("house_calls_used", 0) >= HOUSE_FREE_CALLS_PER_USER:
        return provider, "", False, "user_limit"

    month_key = datetime.now().strftime("%Y-%m")
    house_usage = await db_get_house_usage(month_key)
    if house_usage["cost_usd"] >= HOUSE_MONTHLY_BUDGET_USD:
        return provider, "", False, "global_budget"

    return HOUSE_AI_PROVIDER, HOUSE_AI_API_KEY, True, None


async def call_ai_for_user(prompt: str, user: dict, json_mode: bool = False) -> str:
    provider, api_key, using_house, blocked_reason = await resolve_ai_credentials(user)
    if not api_key:
        if blocked_reason == "user_limit":
            raise HTTPException(
                400, f"体验用的公共额度({HOUSE_FREE_CALLS_PER_USER} 次)已经用完了，去设置里填一个你自己的 AI Key 才能继续用"
            )
        if blocked_reason == "global_budget":
            raise HTTPException(400, "公共体验额度这个月已经用满了，去设置里填一个你自己的 AI Key 才能继续用")
        raise HTTPException(400, "还没有配置 AI API Key，先去设置里填一下")

    relay_base_url = "" if using_house else (user.get("ai_relay_base_url") or "")
    relay_model = "" if using_house else (user.get("ai_relay_model") or "")
    text, cost = await call_ai(
        prompt, provider, api_key, user["id"], json_mode=json_mode,
        relay_base_url=relay_base_url, relay_model=relay_model,
    )

    if using_house:
        await db_increment_house_calls_used(user["id"])
        month_key = datetime.now().strftime("%Y-%m")
        await db_record_house_usage(month_key, cost)

    return text


def build_word_prompt(word: str, context: str, learning_language: str, explain_language: str) -> str:
    learn_label = LANGUAGE_LABELS.get(learning_language, learning_language)
    explain_label = LANGUAGE_LABELS.get(explain_language, explain_language)
    return (
        f"你是一个{learn_label}学习助手。用户正在学习{learn_label}，阅读材料时选中了下面这个词，正在积累生词。\n"
        f"单词/词语：{word}\n"
        f"该词所在的例句：{context or '（无）'}\n\n"
        f"请结合例句的语境，以 JSON 格式返回，包含以下字段，不要输出任何多余文字：\n"
        f'{{"chinese_meaning": "这个词在该语境下的准确释义，用{explain_label}表达，简洁，不超过15个字", '
        f'"ipa": "这个词的注音标记(比如{learn_label}有对应的音标/拼音/罗马音等系统就给出，不带斜杠符号；如果这门语言没有这类概念就留空字符串)", '
        f'"pos": "词性缩写，如 n. / v. / adj. / adv. / prep. 等；如果这门语言没有对应概念就留空字符串"}}'
    )


def build_passage_prompt(text: str, learning_language: str, explain_language: str) -> str:
    learn_label = LANGUAGE_LABELS.get(learning_language, learning_language)
    explain_label = LANGUAGE_LABELS.get(explain_language, explain_language)
    return (
        f"你是一个{learn_label}学习助手。用户正在学习{learn_label}，阅读材料时选中了下面这句话，觉得理解起来有难度。\n"
        f"选中内容：{text}\n\n"
        f"请用简洁的{explain_label}回答，包含：\n"
        "1) 这句话的整体意思\n"
        "2) 语法结构拆解或值得注意的表达方式（如果有难点的话）\n"
        "直接给出结果，不要客套话，不要重复原文。"
    )


class AnalyzeRequest(BaseModel):
    text: str
    context: str = ""
    mode: str  # "word" | "passage"
    learning_language: str = "en"


class AnalyzeResponse(BaseModel):
    mode: str
    explanation: str = ""
    chinese_meaning: Optional[str] = None
    ipa: Optional[str] = None
    pos: Optional[str] = None


@app.post("/api/analyze", response_model=AnalyzeResponse)
@limiter.limit("30/minute")
async def analyze_selection(request: Request, req: AnalyzeRequest, user: dict = Depends(get_current_user)):
    explain_language = resolve_explain_language(user)
    if req.mode == "word":
        prompt = build_word_prompt(req.text, req.context, req.learning_language, explain_language)
        raw = await call_ai_for_user(prompt, user, json_mode=True)
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError:
            parsed = {}
        return AnalyzeResponse(
            mode="word",
            chinese_meaning=parsed.get("chinese_meaning", ""),
            ipa=parsed.get("ipa", ""),
            pos=parsed.get("pos", ""),
        )

    prompt = build_passage_prompt(req.text, req.learning_language, explain_language)
    explanation = await call_ai_for_user(prompt, user)
    return AnalyzeResponse(mode="passage", explanation=explanation)


# ---------- 渐进沉浸阅读模式(把母语文章里挑一些实义词换成目标学习语言) ----------
#
# 按文章的母语(source_language，由后端从文章内容自动检测——不能用文档的 learning_language，
# 那是"这篇文章在学什么语言"，跟"文章本身是什么语言写的"是两回事)
# 分发到不同的分词实现，每个实现的输出形状必须一致：给一段文本，返回一份"实义候选词"字符串列表
# (按出现顺序去重)。这份候选词表有两个用途：(1) 驱动替换比例的斜坡计算；(2) 约束 AI 只能从
# 候选词里选词替换——AI 返回结果后会逐个校验 original_word 是否在对应段落的候选集合里，不在就丢弃。
# 所以"不替换虚词/不替换专有名词/不替换数字"这条规则是在这一层本地强制生效的，不依赖 AI 听话。
#
# 新增一种母语的支持，只需要写一个同接口的 tokenize_xx() 函数，在 IMMERSION_TOKENIZERS 里加一行，
# 不用改这段代码之外的任何替换/插入逻辑。未登记的语言默认退化到英语规则分词。

# jieba 词性标注里，这几个前缀算"实义内容词"(名词/动词/形容词/副词及其派生形式如 vn/an/ad)，
# 才有资格被替换；介词/连词/代词/助词等虚词天然不在这几个前缀里，不用额外排除。
_ZH_CONTENT_POS_PREFIXES = ("n", "v", "a", "d")
# 默认排除专有名词(nr人名/ns地名/nt机构名/nz其它专名)和数字类(m数字/q量词)。
_ZH_EXCLUDED_POS = {"nr", "ns", "nt", "nz", "m", "q"}
# 这几个是语素/成分标记，不是完整的词，不适合单独替换。
_ZH_NON_WORD_POS = {"ng", "vg", "ag", "dg"}


def tokenize_zh(paragraph: str, exclude_proper_nouns: bool) -> list:
    """中文：jieba 分词 + 词性标注筛实义词。"""
    seen = set()
    candidates = []
    for word, flag in pseg.cut(paragraph):
        word = word.strip()
        if not word or word in seen or flag in _ZH_NON_WORD_POS:
            continue
        if exclude_proper_nouns and flag in _ZH_EXCLUDED_POS:
            continue
        if flag.startswith(_ZH_CONTENT_POS_PREFIXES):
            seen.add(word)
            candidates.append(word)
    return candidates


# 英语天然靠空格分词，不需要引入 spaCy/NLTK 这类分词库：按单词字符切 token，用一份静态虚词表
# 排除冠词/介词/连词/代词/助动词等功能词；数字天然不在字母正则里，不用额外处理；
# "排除专有名词"退化成"首字母大写且不是段首第一个词"的启发式，不完美但足够用于这个辅助功能。
_EN_WORD_RE = re.compile(r"[A-Za-z]+(?:'[A-Za-z]+)?")
_EN_STOPWORDS = {
    "a", "an", "the", "and", "or", "but", "if", "then", "than", "so", "because", "as", "of", "in",
    "on", "at", "by", "for", "with", "about", "against", "between", "into", "through", "during",
    "before", "after", "above", "below", "to", "from", "up", "down", "out", "off", "over", "under",
    "again", "further", "once", "here", "there", "when", "where", "why", "how", "all", "any", "both",
    "each", "few", "more", "most", "other", "some", "such", "no", "nor", "not", "only", "own", "same",
    "i", "me", "my", "myself", "we", "our", "ours", "ourselves", "you", "your", "yours", "yourself",
    "yourselves", "he", "him", "his", "himself", "she", "her", "hers", "herself", "it", "its",
    "itself", "they", "them", "their", "theirs", "themselves", "what", "which", "who", "whom",
    "this", "that", "these", "those", "am", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "having", "do", "does", "did", "doing", "will", "would", "shall", "should",
    "can", "could", "may", "might", "must", "ought", "very", "just", "also", "too",
}


def tokenize_en(paragraph: str, exclude_proper_nouns: bool) -> list:
    """英语：规则分词，无需第三方分词库。"""
    seen = set()
    candidates = []
    for i, m in enumerate(_EN_WORD_RE.finditer(paragraph)):
        word = m.group(0)
        lower = word.lower()
        if len(word) < 2 or lower in seen or lower in _EN_STOPWORDS:
            continue
        if exclude_proper_nouns and word[0].isupper() and i > 0:
            continue
        seen.add(lower)
        candidates.append(word)
    return candidates


# 韩语：soynlp 是无监督统计分词，需要在真实语料上训练才能获得可靠的词边界，不适合对单个、
# 零散的段落即时分词，所以这里改用和英语同思路的规则实现——按谚文字符切出어절(eojeol)，
# 剥掉常见的助词/词尾后缀，再用一份停用词表过滤代词/指示词等功能词。韩语没有大小写信号，
# 无法像英语那样启发式识别专有名词，exclude_proper_nouns 对韩语暂不生效。
_KO_TOKEN_RE = re.compile(r"[가-힣]+")
_KO_PARTICLE_SUFFIXES = sorted([
    "에게서", "으로부터", "이라고", "이라도", "에서", "에게", "한테", "부터", "까지", "이나",
    "이며", "이랑", "만큼", "밖에", "처럼", "보다", "으로", "라고", "라도", "이", "가", "은",
    "는", "을", "를", "의", "에", "와", "과", "도", "만", "나", "며", "랑", "로", "조차", "마저",
], key=len, reverse=True)
_KO_STOPWORDS = {
    "그", "이", "저", "것", "수", "때", "등", "및", "또는", "그리고", "그러나", "하지만", "그런데",
    "즉", "따라서", "그래서", "때문", "대해", "대한", "통해", "위해", "의해", "나", "너", "우리",
    "저희", "당신", "여기", "거기", "저기", "어디", "언제", "무엇", "누구", "어떻게", "왜", "매우",
    "정말", "너무", "아주", "좀", "조금", "항상", "자주", "가끔", "모든", "여러", "각", "이런",
    "그런", "저런", "합니다", "있다", "없다", "하다",
}


def tokenize_ko(paragraph: str, exclude_proper_nouns: bool) -> list:
    """韩语：规则分词(어절 + 助词剥离)，不依赖 soynlp/KoNLPy。"""
    seen = set()
    candidates = []
    for m in _KO_TOKEN_RE.finditer(paragraph):
        stem = m.group(0)
        for suf in _KO_PARTICLE_SUFFIXES:
            if stem.endswith(suf) and len(stem) > len(suf):
                stem = stem[: -len(suf)]
                break
        if len(stem) < 2 or stem in seen or stem in _KO_STOPWORDS:
            continue
        seen.add(stem)
        candidates.append(stem)
    return candidates


# 日语：janome 是纯 Python 实现、自带小型词典，pip 装完即可用，不依赖 MeCab/Java，
# 词性标注格式形如 "名詞,固有名詞,人名,一般"，用逗号切开取大类/细类。
_JA_CONTENT_POS_MAJOR = ("名詞", "動詞", "形容詞", "副詞")
_JA_NON_WORD_POS_DETAIL = {"非自立", "接尾", "代名詞", "数"}
# サ変接続名詞(発展/投資等)+する/し 这种复合动词，janome 会把する/し单独切出来，标成"動詞,自立"，
# 但单个假名的し/い/き这类几乎都是这种被拆开的语法残片，不是有意义的独立词，直接按长度过滤掉。
_JA_HIRAGANA_RE = re.compile(r"^[ぁ-ん]+$")
_janome_tokenizer = None


def _get_janome_tokenizer() -> JanomeTokenizer:
    global _janome_tokenizer
    if _janome_tokenizer is None:
        _janome_tokenizer = JanomeTokenizer()
    return _janome_tokenizer


def tokenize_ja(paragraph: str, exclude_proper_nouns: bool) -> list:
    """日语：janome 分词 + 词性标注筛实义词。"""
    seen = set()
    candidates = []
    for token in _get_janome_tokenizer().tokenize(paragraph):
        word = token.surface.strip()
        if not word or word in seen:
            continue
        if len(word) == 1 and _JA_HIRAGANA_RE.match(word):
            continue
        pos_parts = token.part_of_speech.split(",")
        major = pos_parts[0]
        detail = pos_parts[1] if len(pos_parts) > 1 else ""
        if detail in _JA_NON_WORD_POS_DETAIL:
            continue
        if exclude_proper_nouns and detail == "固有名詞":
            continue
        if major.startswith(_JA_CONTENT_POS_MAJOR):
            seen.add(word)
            candidates.append(word)
    return candidates


# 母语语言 -> 分词器分发表，仿照 LANGUAGE_SOURCES 的"按语言 key 查配置"模式。
IMMERSION_TOKENIZERS = {
    "zh": {"tokenizer": "jieba"},
    "en": {"tokenizer": "en_rule"},
    "ko": {"tokenizer": "ko_rule"},
    "ja": {"tokenizer": "janome"},
}
_IMMERSION_TOKENIZER_FUNCS = {
    "jieba": tokenize_zh,
    "en_rule": tokenize_en,
    "ko_rule": tokenize_ko,
    "janome": tokenize_ja,
}

IMMERSION_PARAGRAPH_CHUNK_SIZE = 7  # 每次 AI 调用最多带几个段落，避免一篇文章打成几十次调用


def extract_immersion_candidates(paragraph: str, source_language: str, exclude_proper_nouns: bool) -> list:
    """给一段母语文本分词，挑出可以被替换的实义词，按出现顺序去重。按 source_language 分发到
    对应的分词实现；未登记的语言默认退化到英语规则分词。"""
    tokenizer_key = IMMERSION_TOKENIZERS.get(source_language, IMMERSION_TOKENIZERS["en"])["tokenizer"]
    return _IMMERSION_TOKENIZER_FUNCS[tokenizer_key](paragraph, exclude_proper_nouns)


# ---------- 生词覆盖率(跟沉浸模式共用底层分词器，但统计口径不同：这里要统计"这篇文章一共多少个
# 词"，所以不筛词性、不去重，跟前端"约 N 词"的字数统计口径保持一致；沉浸模式那套 tokenize_xx()
# 是特意筛过的实义词候选集，直接拿来当分母会把覆盖率算得偏高，所以中/日/韩单独写这三个"数全部"版本) ----------

def tokenize_coverage_zh(paragraph: str) -> list:
    return [w for w in jieba.cut(paragraph) if w.strip() and re.search(r"[\w一-鿿]", w)]


def tokenize_coverage_ja(paragraph: str) -> list:
    tokens = []
    for token in _get_janome_tokenizer().tokenize(paragraph):
        word = token.surface.strip()
        if not word or (len(word) == 1 and _JA_HIRAGANA_RE.match(word)):
            continue
        tokens.append(word)
    return tokens


def tokenize_coverage_ko(paragraph: str) -> list:
    tokens = []
    for m in _KO_TOKEN_RE.finditer(paragraph):
        stem = m.group(0)
        for suf in _KO_PARTICLE_SUFFIXES:
            if stem.endswith(suf) and len(stem) > len(suf):
                stem = stem[: -len(suf)]
                break
        if stem:
            tokens.append(stem)
    return tokens


COVERAGE_TOKENIZERS = {
    "zh": tokenize_coverage_zh,
    "ja": tokenize_coverage_ja,
    "ko": tokenize_coverage_ko,
}


def compute_immersion_replace_counts(candidate_counts: list, start_pct: float, end_pct: float) -> list:
    """candidate_counts[i] = 第 i 段候选词数量；按段落顺序在 start_pct~end_pct 之间线性递增，
    返回每段该替换几个词。"""
    n = len(candidate_counts)
    counts = []
    for i, n_candidates in enumerate(candidate_counts):
        pct = start_pct + (end_pct - start_pct) * i / max(1, n - 1)
        counts.append(min(n_candidates, round(pct / 100 * n_candidates)))
    return counts


def build_immersion_prompt(paragraph_group: list, learning_language: str, explain_language: str,
                            known_vocab: list, source_priority: str) -> str:
    """paragraph_group: [{"index", "text", "candidates", "replace_count"}, ...]"""
    lang_label = LANGUAGE_LABELS.get(learning_language, learning_language)
    explain_label = LANGUAGE_LABELS.get(explain_language, explain_language)

    listing = "\n\n".join(
        f'段落{p["index"]}(需要替换{p["replace_count"]}个词): "{p["text"]}"\n'
        f'候选词(只能从这些词里选，不要选列表外的词): {"、".join(p["candidates"])}'
        for p in paragraph_group
    )

    if source_priority == "vocab" and known_vocab:
        priority_hint = f"优先选下面这个用户正在学的{lang_label}生词表里已有对应意思的词：{', '.join(known_vocab)}；生词表里没有对应意思的，再从候选词里挑常见词翻译"
    else:
        priority_hint = f"优先挑常见、高频的{lang_label}词汇，不要选生僻词"

    return (
        f"你在帮一个正在学{lang_label}的用户做「沉浸式阅读」——把母语文章里的一部分实义词换成对应的"
        f"{lang_label}词，读起来大部分还是母语，但穿插一些{lang_label}词帮助学习。\n"
        f"{priority_hint}。\n"
        "每个段落后面给了一份「候选词」列表(已经按词性筛过，都是名词/动词/形容词/副词这类实义词)，"
        "请从每段候选词列表里，挑出题目要求数量的词，翻译成对应的" + lang_label + "词。"
        "只能选候选词列表里原样出现的词，不要自己造词或选列表之外的词。\n\n"
        f"{listing}\n\n"
        '请以 JSON 格式返回：{"paragraphs": [{"index": 段落编号(数字), "substitutions": '
        f'[{{"original_word": "母语原词(必须是候选词列表里的原词)", "target_word": "翻译成的{lang_label}词", '
        f'"ipa": "该{lang_label}词的注音(音标/拼音/罗马音等，没有就留空字符串)", '
        '"pos": "词性缩写，如 n./v./adj./adv.，没有就留空字符串"}]}]}\n'
        "只返回这个 JSON，不要有其他文字。"
    )


class ImmersionPlanRequest(BaseModel):
    content: str
    source_language: str = ""  # 文章本身的语言；留空(默认)由后端从内容自动检测，不要传文档的 learning_language
    learning_language: str = "en"
    start_ratio: float = 10  # 百分比，10 表示 10%
    end_ratio: float = 30
    exclude_proper_nouns: bool = True
    source_priority: str = "vocab"  # "vocab" | "frequency"


class ImmersionSubstitution(BaseModel):
    original_word: str
    target_word: str
    ipa: str = ""
    pos: str = ""


class ImmersionParagraphPlan(BaseModel):
    index: int
    substitutions: list[ImmersionSubstitution]


class ImmersionPlanResponse(BaseModel):
    paragraphs: list[ImmersionParagraphPlan]


@app.post("/api/immersion/plan", response_model=ImmersionPlanResponse)
@limiter.limit("10/minute")
async def get_immersion_plan(request: Request, req: ImmersionPlanRequest, user: dict = Depends(get_current_user)):
    if req.learning_language not in LANGUAGE_LABELS:
        raise HTTPException(400, f"不支持的目标语言: {req.learning_language}")
    if req.source_priority not in IMMERSION_SOURCE_PRIORITY_CHOICES:
        raise HTTPException(400, "不支持的选词策略")

    # 沉浸模式一篇文章要拆好几次 AI 调用，公共体验额度经不起这么用，只对配了自己 key 的用户开放。
    own_key = user.get("ai_api_keys", {}).get(user.get("ai_provider", "deepseek"), "")
    if not own_key:
        raise HTTPException(400, "沉浸阅读模式一次要消耗好几次 AI 调用，暂时只对配置了自己 AI Key 的用户开放，去设置里填一个吧")

    # 跟前端 renderTextDocument() 的 content.split(/\n+/) + 去空段 保持段落索引完全对齐，
    # 这里是等价写法：按换行切开，过滤掉空段。
    paragraphs = [p.strip() for p in re.split(r"\n+", req.content) if p.strip()]
    if not paragraphs:
        return ImmersionPlanResponse(paragraphs=[])

    # 文章母语从内容自动检测(字符区间判断，中/日/韩等直接可辨；拉丁语系统一按英语规则分词，
    # 对法/西/德也够用——同样是空格分词+首字母大写判专名)。请求里显式传了就用传的。
    source_language = req.source_language.strip() or detect_learning_language(req.content)
    candidates_per_paragraph = [
        extract_immersion_candidates(p, source_language, req.exclude_proper_nouns) for p in paragraphs
    ]
    replace_counts = compute_immersion_replace_counts(
        [len(c) for c in candidates_per_paragraph], req.start_ratio, req.end_ratio
    )

    known_vocab = []
    if req.source_priority == "vocab":
        vocab_rows = await db_list_vocab(user["id"])
        known_vocab = [
            v["word"] for v in vocab_rows
            if v.get("learning_language") == req.learning_language and v.get("word")
        ][:80]  # 避免生词表太长把 prompt 撑爆

    explain_language = resolve_explain_language(user)

    groups = []
    current_group = []
    for i, (para_text, candidates, count) in enumerate(zip(paragraphs, candidates_per_paragraph, replace_counts)):
        if count <= 0 or not candidates:
            continue
        current_group.append({"index": i, "text": para_text, "candidates": candidates, "replace_count": count})
        if len(current_group) >= IMMERSION_PARAGRAPH_CHUNK_SIZE:
            groups.append(current_group)
            current_group = []
    if current_group:
        groups.append(current_group)

    plan_paragraphs = {}
    for group in groups:
        prompt = build_immersion_prompt(group, req.learning_language, explain_language, known_vocab, req.source_priority)
        raw = await call_ai_for_user(prompt, user, json_mode=True)
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError:
            continue
        for p in parsed.get("paragraphs", []):
            idx = p.get("index")
            if not isinstance(idx, int) or not (0 <= idx < len(candidates_per_paragraph)):
                continue
            valid_candidates = set(candidates_per_paragraph[idx])
            subs = []
            for s in p.get("substitutions", []):
                original = (s.get("original_word") or "").strip()
                target = (s.get("target_word") or "").strip()
                # 只信任真的来自候选词列表的词，防止 AI 编出文章里没有的词导致前端替换失败/找不到位置。
                if not original or not target or original not in valid_candidates:
                    continue
                subs.append(ImmersionSubstitution(
                    original_word=original, target_word=target,
                    ipa=s.get("ipa", ""), pos=s.get("pos", ""),
                ))
            if subs:
                plan_paragraphs[idx] = subs

    return ImmersionPlanResponse(
        paragraphs=[
            ImmersionParagraphPlan(index=idx, substitutions=subs)
            for idx, subs in sorted(plan_paragraphs.items())
        ]
    )


# ---------- AI 文章推荐(按学习语言匹配新闻源，拉标题 + AI 按难度/话题筛选打标签) ----------

# 每种学习语言对应的推荐新闻源，按顺序尝试(第一个是首选源，后面的当兜底/补充)。
# 除了 en 之外的这几个源地址是按各家媒体一贯的 RSS 惯例整理的，这个沙箱环境出站网络受限，
# 没能逐一实测连通性——上线后如果发现某个源失效或改版了，把对应条目删掉/换新地址即可，
# fetch_headlines() 已经做了"单个源抓取失败不影响其它源、也不会导致整个接口报错"的容错。
# de/fr/ja 原本想接各家的"简易/学习者"专区(DW Deutsch Lernen、RFI Français Facile、
# NHK NEWS EASY)，但这几个专区要么没有公开 RSS，要么用的是非标准接口，上线后确认连不通；
# 换成了这几家媒体本身有文档、其它项目也验证过的常规新闻 RSS——内容不是专门简化过的，
# 但 AI 推荐时本来就会按用户水平挑难度适中的文章，不是完全没有筛选。
LANGUAGE_SOURCES = {
    "en": [
        {"name": "BBC", "url": "https://feeds.bbci.co.uk/news/rss.xml"},
        {"name": "BBC", "url": "https://feeds.bbci.co.uk/news/technology/rss.xml"},
        {"name": "BBC", "url": "https://feeds.bbci.co.uk/news/business/rss.xml"},
        {"name": "BBC", "url": "https://feeds.bbci.co.uk/news/science_and_environment/rss.xml"},
        {"name": "Guardian", "url": "https://www.theguardian.com/world/rss"},
        {"name": "Guardian", "url": "https://www.theguardian.com/culture/rss"},
        {"name": "Guardian", "url": "https://www.theguardian.com/science/rss"},
    ],
    "ko": [
        {"name": "연합뉴스", "url": "https://www.yna.co.kr/rss/news.xml"},
        {"name": "KBS뉴스", "url": "https://world.kbs.co.kr/rss/rss_news.htm?lang=k"},
    ],
    "ja": [
        {"name": "NHK", "url": "https://www3.nhk.or.jp/rss/news/cat0.xml"},
    ],
    "fr": [
        {"name": "RFI", "url": "https://www.rfi.fr/fr/rss"},
    ],
    "de": [
        {"name": "DW", "url": "https://rss.dw.com/rdf/rss-de-top"},
        {"name": "DW", "url": "https://rss.dw.com/rdf/rss-de-all"},
    ],
    "es": [
        {"name": "BBC Mundo", "url": "https://feeds.bbci.co.uk/mundo/rss.xml"},
    ],
}
DEFAULT_RECOMMEND_LANGUAGE = "en"

# 只有 en 是根据这个具体用户的真实水平写的；其它语言用一个通用的中等水平描述。
RECOMMEND_LEVEL_DESC = {
    "en": "大学英语四级已通过，六级考了480分，属于中等偏下的中高级学习者（约B1-B2水平）",
}
DEFAULT_RECOMMEND_LEVEL_DESC = "中等水平的语言学习者（约B1-B2水平），能读懂大意但生词量还有限"

_recommend_cache = {}  # (user_id, learning_language) -> {"data": [...], "ts": datetime}
RECOMMEND_CACHE_SECONDS = 3 * 60 * 60


def _strip_html(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text or "").strip()


def _fetch_rss_headlines(source_name: str, feed_url: str) -> list:
    items = []
    try:
        parsed = feedparser.parse(feed_url)
        for entry in parsed.entries[:8]:
            title = entry.get("title", "").strip()
            url = entry.get("link", "").strip()
            if not title or not url:
                continue
            items.append({
                "source": source_name,
                "title": title,
                "summary": _strip_html(entry.get("summary", ""))[:220],
                "url": url,
            })
    except Exception:
        pass
    return items


def fetch_headlines(learning_language: str) -> list:
    sources = LANGUAGE_SOURCES.get(learning_language) or LANGUAGE_SOURCES[DEFAULT_RECOMMEND_LANGUAGE]
    items = []
    for source in sources:
        items.extend(_fetch_rss_headlines(source["name"], source["url"]))
    return items


def build_recommend_prompt(items: list, learning_language: str, explain_language: str) -> str:
    lang_label = LANGUAGE_LABELS.get(learning_language, learning_language)
    explain_label = LANGUAGE_LABELS.get(explain_language, explain_language)
    level_desc = RECOMMEND_LEVEL_DESC.get(learning_language, DEFAULT_RECOMMEND_LEVEL_DESC)
    listing = "\n".join(f"{i}. [{it['source']}] {it['title']} — {it['summary']}" for i, it in enumerate(items))
    return (
        f"你是一个{lang_label}学习内容推荐助手。用户正在学习{lang_label}，水平：{level_desc}。\n"
        f"下面是一批当天的{lang_label}新闻标题和摘要，请帮用户从中挑出 6-8 篇适合精读积累的文章。\n"
        "挑选标准：\n"
        "1) 难度要适中——不要挑长难句堆砌、专业术语密集的深度调查或学术性文章，也不要挑过短的快讯简报\n"
        "2) 话题尽量多样化，覆盖不同领域，不要挑到好几篇话题重复的\n"
        "3) 优先挑叙事性强、可读性好、有完整信息量的文章\n\n"
        f"{listing}\n\n"
        f"请以 JSON 格式返回，其中 tags / difficulty / reason 这三个字段的内容都要用{explain_label}表达"
        f"(不要用{lang_label}或其它语言)：\n"
        '{"picks": [{"index": 原列表序号(数字), '
        f'"tags": ["1到2个{explain_label}话题标签"], '
        f'"difficulty": "用{explain_label}表达的难度，大致对应中等偏易/中等/偏难三档中的一档", '
        f'"reason": "一句话说明推荐理由，用{explain_label}表达，不超过30字"}}]}}\n'
        "只返回这个 JSON，不要有其他文字。"
    )


@app.get("/api/recommendations")
async def get_recommendations(
    refresh: bool = False,
    learning_language: str = "en",
    user: dict = Depends(get_current_user),
):
    if learning_language not in LANGUAGE_SOURCES:
        lang_label = LANGUAGE_LABELS.get(learning_language, learning_language)
        supported = "、".join(LANGUAGE_LABELS.get(code, code) for code in LANGUAGE_SOURCES)
        raise HTTPException(400, f"暂时还没有给「{lang_label}」配置推荐新闻源，目前支持：{supported}")

    lang = learning_language
    explain_language = resolve_explain_language(user)
    cache_key = (user["id"], lang, explain_language)
    now = datetime.now(timezone.utc)
    cache_entry = _recommend_cache.get(cache_key)
    if not refresh and cache_entry:
        age = (now - cache_entry["ts"]).total_seconds()
        if age < RECOMMEND_CACHE_SECONDS:
            return cache_entry["data"]

    items = await asyncio.to_thread(fetch_headlines, lang)
    if not items:
        source_names = dict.fromkeys(s["name"] for s in LANGUAGE_SOURCES.get(lang, []))
        raise HTTPException(502, f"{' / '.join(source_names)} 暂时无法访问，请稍后重试")

    raw = await call_ai_for_user(build_recommend_prompt(items, lang, explain_language), user, json_mode=True)
    try:
        picks = json.loads(raw).get("picks", [])
    except json.JSONDecodeError:
        picks = []

    results = []
    seen_urls = set()
    for p in picks:
        idx = p.get("index")
        if not isinstance(idx, int) or not (0 <= idx < len(items)):
            continue
        item = items[idx]
        if item["url"] in seen_urls:
            continue
        seen_urls.add(item["url"])
        results.append({
            **item,
            "tags": p.get("tags", []),
            "difficulty": p.get("difficulty", ""),
            "reason": p.get("reason", ""),
        })

    _recommend_cache[cache_key] = {"data": results, "ts": now}
    return results


# ---------- 设置(AI 服务商 + key、Google Sheets 同步开关) ----------

def mask_key(key: str) -> str:
    if not key:
        return ""
    if len(key) <= 8:
        return "*" * len(key)
    return key[:4] + "..." + key[-4:]


@app.get("/api/settings")
async def get_settings(user: dict = Depends(get_current_user)):
    keys = user.get("ai_api_keys", {})
    return {
        "name": user.get("name", ""),
        "is_owner": user.get("is_owner", False),
        "has_password": bool(user.get("password_hash")),
        "has_google": bool(user.get("google_id")),
        "google_email": user.get("google_email") or "",
        "ai_provider": user.get("ai_provider", "deepseek"),
        "ai_key_status": {
            k: {"has_key": bool(keys.get(k)), "masked": mask_key(keys.get(k, ""))}
            for k in PROVIDER_CONFIG
        },
        "sheets_sync_enabled": user.get("sheets_sync_enabled", False),
        "providers": [{"value": k, "label": v["label"]} for k, v in PROVIDER_CONFIG.items()],
        "house_trial_enabled": bool(HOUSE_AI_API_KEY),
        "house_calls_used": user.get("house_calls_used", 0),
        "house_calls_total": HOUSE_FREE_CALLS_PER_USER,
        "ui_language": user.get("ui_language", "en"),
        "ui_languages": UI_LANGUAGES,
        "explain_language": user.get("explain_language", "auto"),
        "explain_language_choices": EXPLAIN_LANGUAGE_CHOICES,
        "ai_relay_base_url": user.get("ai_relay_base_url", ""),
        "ai_relay_model": user.get("ai_relay_model", ""),
        "immersion_target_language": user.get("immersion_target_language", "follow"),
        "immersion_source_priority": user.get("immersion_source_priority", "vocab"),
        "immersion_exclude_proper_nouns": user.get("immersion_exclude_proper_nouns", True),
    }


IMMERSION_SOURCE_PRIORITY_CHOICES = ["vocab", "frequency"]


class SettingsRequest(BaseModel):
    ai_provider: str
    ai_api_key: str = ""  # 留空表示不修改这个服务商已保存的 key
    sheets_sync_enabled: bool = False
    ui_language: str = "en"
    explain_language: str = "auto"
    ai_relay_base_url: str = ""  # 中转站地址，配了之后四个服务商都走这个地址，留空用官方地址
    ai_relay_model: str = ""  # 中转站用的模型名，留空则用各服务商的官方默认模型名
    immersion_target_language: str = "follow"  # 沉浸模式目标语言，"follow" 表示跟随文章的学习语言
    immersion_source_priority: str = "vocab"  # "vocab"=生词本优先 / "frequency"=高频词优先
    immersion_exclude_proper_nouns: bool = True


@app.post("/api/settings")
async def update_settings(req: SettingsRequest, user: dict = Depends(get_current_user)):
    if req.ai_provider not in PROVIDER_CONFIG:
        raise HTTPException(400, "不支持的 AI 服务商")
    if req.ui_language not in UI_LANGUAGES:
        raise HTTPException(400, "不支持的界面语言")
    if req.explain_language not in EXPLAIN_LANGUAGE_CHOICES:
        raise HTTPException(400, "不支持的 AI 讲解语言")
    if req.immersion_target_language not in ({"follow"} | set(LANGUAGE_SOURCES.keys())):
        raise HTTPException(400, "不支持的沉浸模式目标语言")
    if req.immersion_source_priority not in IMMERSION_SOURCE_PRIORITY_CHOICES:
        raise HTTPException(400, "不支持的沉浸模式选词策略")

    relay_base_url = req.ai_relay_base_url.strip()
    if relay_base_url and not (relay_base_url.startswith("http://") or relay_base_url.startswith("https://")):
        raise HTTPException(400, "中转站地址得是 http:// 或 https:// 开头的完整地址")

    keys = dict(user.get("ai_api_keys", {}))
    if req.ai_api_key:
        keys[req.ai_provider] = req.ai_api_key
    update_fields = {
        "ai_provider": req.ai_provider,
        "ai_api_keys": keys,
        "ui_language": req.ui_language,
        "explain_language": req.explain_language,
        "ai_relay_base_url": relay_base_url,
        "ai_relay_model": req.ai_relay_model.strip(),
        "immersion_target_language": req.immersion_target_language,
        "immersion_source_priority": req.immersion_source_priority,
        "immersion_exclude_proper_nouns": req.immersion_exclude_proper_nouns,
    }
    if user.get("is_owner"):
        update_fields["sheets_sync_enabled"] = req.sheets_sync_enabled
    await db_update_user_fields(user["id"], **update_fields)
    return {"ok": True}


# ---------- 保存生词 / 句子笔记(本地 + 可选 Google Sheet 同步，仅限主账号) ----------

async def push_to_sheet(payload: dict) -> bool:
    if not SHEETS_WEBHOOK_URL:
        return False
    try:
        async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
            resp = await client.post(SHEETS_WEBHOOK_URL, json={**payload, "token": SHEETS_TOKEN})
        return resp.status_code == 200
    except Exception:
        return False


class SaveRequest(BaseModel):
    mode: str  # "word" | "passage"
    text: str
    context: str = ""
    explanation: str = ""
    chinese_meaning: str = ""
    ipa: str = ""
    pos: str = ""
    source_doc: str = ""
    learning_language: str = "en"


@app.post("/api/save")
async def save_entry(req: SaveRequest, user: dict = Depends(get_current_user)):
    today = datetime.now().strftime("%Y-%m-%d")
    now_iso = datetime.now(timezone.utc).isoformat()
    should_sync = bool(user.get("is_owner")) and bool(user.get("sheets_sync_enabled"))
    explain_language = resolve_explain_language(user)

    if req.mode == "word":
        existing = await db_find_vocab_by_word(user["id"], req.text)
        if existing:
            return {"record": existing, "sheet_synced": False, "duplicate": True}

        record = {
            "id": uuid.uuid4().hex[:12],
            "user_id": user["id"],
            "word": req.text,
            "sentence": req.context,
            "chinese_meaning": req.chinese_meaning,
            "ipa": req.ipa,
            "pos": req.pos,
            "source_doc": req.source_doc,
            "date": today,
            "learning_language": req.learning_language,
            "explain_language": explain_language,
            "added_at": now_iso,
        }
        await db_create_vocab(record)
        synced = False
        if should_sync:
            synced = await push_to_sheet({
                "type": "word",
                "word": req.text,
                "sentence": req.context,
                "chinese_meaning": req.chinese_meaning,
                "ipa": req.ipa,
                "pos": req.pos,
                "source": req.source_doc,
                "date": today,
            })
        return {"record": record, "sheet_synced": synced, "duplicate": False}
    else:
        record = {
            "id": uuid.uuid4().hex[:12],
            "user_id": user["id"],
            "sentence": req.text,
            "analysis": req.explanation,
            "source_doc": req.source_doc,
            "date": today,
            "learning_language": req.learning_language,
            "explain_language": explain_language,
            "added_at": now_iso,
        }
        await db_create_sentence_note(record)
        synced = False
        if should_sync:
            synced = await push_to_sheet({
                "type": "sentence",
                "sentence": req.text,
                "analysis": req.explanation,
                "source": req.source_doc,
                "date": today,
            })
        return {"record": record, "sheet_synced": synced, "duplicate": False}


@app.get("/api/vocab")
async def list_vocab(user: dict = Depends(get_current_user)):
    return await db_list_vocab(user["id"])


@app.get("/api/sentence_notes")
async def list_sentence_notes(user: dict = Depends(get_current_user)):
    return await db_list_sentence_notes(user["id"])


@app.delete("/api/vocab/{item_id}")
async def delete_vocab(item_id: str, user: dict = Depends(get_current_user)):
    target = await db_get_vocab(item_id)
    if not target or target.get("user_id") != user["id"]:
        raise HTTPException(404, "没找到这条生词记录")
    await db_delete_vocab(item_id)
    return {"ok": True}


REVIEW_BATCH_SIZE = 20
REVIEW_BATCH_SIZE_MAX = 50


@app.get("/api/review/due-counts")
async def review_due_counts(user: dict = Depends(get_current_user)):
    return await db_count_due_vocab_by_language(user["id"])


@app.get("/api/review/queue")
async def review_queue(learning_language: str, limit: int = REVIEW_BATCH_SIZE, user: dict = Depends(get_current_user)):
    limit = max(1, min(limit, REVIEW_BATCH_SIZE_MAX))
    return await db_list_due_vocab(user["id"], learning_language, limit)


class ReviewMarkRequest(BaseModel):
    item_id: str
    result: str  # "know" | "dont_know"


@app.post("/api/review/mark")
async def review_mark(req: ReviewMarkRequest, user: dict = Depends(get_current_user)):
    if req.result not in ("know", "dont_know"):
        raise HTTPException(400, "不支持的复习结果")
    target = await db_get_vocab(req.item_id)
    if not target or target.get("user_id") != user["id"]:
        raise HTTPException(404, "没找到这条生词记录")
    new_level, next_review_at = apply_leitner_result(target.get("srs_level", 0), req.result)
    await db_update_vocab_fields(
        req.item_id, srs_level=new_level, next_review_at=next_review_at,
        last_reviewed_at=datetime.now(timezone.utc),
    )
    return {"record": await db_get_vocab(req.item_id)}


@app.delete("/api/sentence_notes/{item_id}")
async def delete_sentence_note(item_id: str, user: dict = Depends(get_current_user)):
    target = await db_get_sentence_note(item_id)
    if not target or target.get("user_id") != user["id"]:
        raise HTTPException(404, "没找到这条句子笔记")
    await db_delete_sentence_note(item_id)
    return {"ok": True}


app.mount("/", StaticFiles(directory=str(FRONTEND_DIR), html=True), name="frontend")
